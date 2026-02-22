#!/usr/bin/env python3
"""Generate professional DOCX report from report-sections.ts data."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import os
import tempfile
from PIL import Image, ImageDraw

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(BASE_DIR, 'public', 'images')

# ============================================================
# 1. Parse TypeScript data
# ============================================================

REPORT_SUMMARY = """기술 혁명은 자동으로 사회를 이롭게 하지 않는다. 지난 230년간 다섯 차례의 기술 혁명에서, 새 기술의 혜택이 사회 전체로 퍼진 황금기는 제도적 재편에 성공한 경우에만 도달했다. 재편에 실패하면 긴장이 격변으로 폭발했다(Perez, 2002).

산업 시대에 이 재편은 대부분 사후적으로 이루어졌다. 시장이 먼저 질서를 만들고, 법이 뒤따랐다. 공장이 돌아가고 기업이 성장한 뒤에야 법인이라는 법적 범주가 만들어졌고, 착취와 독점이 심화된 뒤에야 노동법과 반독점법이 따라왔다. 이 경로는 작동했지만, 대가가 컸다. 역설적이지만, 격변의 경로가 때로 더 높은 황금기에 도달하기도 했다 \u2014 전후 황금기(1945~75)가 그 증거다. 그러나 그 대가는 수천만 명의 죽음이었다.

본 연구는 AI 시대에 이 사후적 경로가 더 이상 안전하지 않다고 주장한다. 산업 시대의 기계는 인간이 끄면 멈췄다. 그래서 "먼저 돌려보고 나중에 고치는" 전략이 가능했다. AI는 점점 더 스스로 판단하고 행동한다. 사후 교정의 비용과 위험이 질적으로 달라진다. 따라서 질서의 일부를 미리 설계할 필요가 있다.

다만, 경제 질서를 위에서 설계하려 한 시도 \u2014 중앙계획경제가 대표적이다 \u2014 는 역사적으로 실패했다. 본 연구가 제안하는 것은 질서 자체의 설계가 아니라 질서가 작동할 조건의 설계다. 경쟁의 틀, 주체의 역할 범위, 책임의 귀속 구조를 미리 설정하되, 그 안에서의 구체적 질서는 시장과 사회에 맡기는 것이다.

이 조건 설계의 경로를 본 연구는 세 단계로 제시한다:"""

SUMMARY_BULLETS = [
    "\u2776 주체 정의 \u2014 AI 에이전트에 대한 법적 범주의 창설",
    "\u2777 관계 정의 \u2014 인간-AI 사이의 권리\u00b7의무\u00b7책임 구조 확립",
    "\u2778 질서 설계 \u2014 경제적\u00b7정치적 질서의 조건 설계 (현실에 맞춰 지속 조정)",
]

# Sections data extracted from TypeScript
SECTIONS = [
    {
        "title": "1. 기술 혁명은 왜 위험한가",
        "subsections": [
            {
                "title": "1.1 황금기는 자동으로 오지 않는다",
                "paragraphs": [
                    "베네수엘라 출신의 기술혁명 경제학자 카를로타 페레즈(2002)는 지난 230년간 다섯 차례의 기술 혁명 \u2014 산업혁명, 증기\u00b7철도, 철강\u00b7전기, 석유\u00b7자동차, 정보\u00b7통신 \u2014 을 분석하여, 놀라울 정도로 일관된 패턴을 발견했다.",
                    "새로운 기술이 등장하면 생산력이 폭발한다. 투기 자본이 몰리고 양극화가 심화된다. 어느 시점에 제도적 전환이 이루어지거나, 이루어지지 않는다. 성공한 경우에만 기술의 혜택이 사회 전체로 퍼지는 황금기에 도달한다.",
                    "이 패턴에서 핵심은 전환점이다. 기술이 발전하면 사회가 저절로 좋아지는 것이 아니다. 230년간 다섯 차례 모두에서, 제도적 재편 없이 황금기에 도달한 사례는 없었다.",
                ],
            },
            {
                "title": "1.2 두 가지 경로",
                "paragraphs": [
                    "그런데 재편이 이루어지는 방식에는 중요한 차이가 있었다.",
                    "대부분의 경우, 재편은 심각한 사회적 비용을 수반했다. 산업혁명 때 긴장은 러다이트 운동과 프랑스 혁명으로 터졌다. 철강\u00b7전기 시대에는 노동 착취와 독점이 수십 년간 방치되었고, 이 긴장이 여러 요인들과 결합하여 1차 세계대전의 토양이 되었다. 석유\u00b7자동차 시대에는 대공황과 2차 세계대전을 거쳐서야 전후 황금기(1945~75)에 도달했다. 제도적 재편이 늦어진 만큼, 기술의 힘과 사회의 통제력 사이의 간극이 벌어졌고, 그 간극이 파괴적 결과의 토양을 만들었다(Perez, 2002).",
                    "그러나 역사가 격변만을 기록한 것은 아니다. 증기\u00b7철도 시대에 영국 의회는 1844~56년에 걸쳐 법인을 일반화했다 \u2014 혁명이 아니라 입법이 전환점을 만든 사례다. 미국의 진보 시대 개혁 \u2014 셔먼법(1890), 연방준비제도(1913), 클레이턴법(1914) \u2014 은 1차 세계대전 이전에 선제적으로 추진되었다. 가장 주목할 것은 스웨덴이다. 살트셰바덴 협약(1938)은 노사가 전쟁이나 혁명 없이 자발적으로 합의한 대타협이었고, 북유럽 복지국가의 토대가 되었다.",
                    "역설적이지만, 격변의 경로가 때로 더 높은 황금기에 도달하기도 했다. 전후 황금기(1945~75)는 인류 역사상 가장 광범위한 번영을 이루었는데, 이는 두 차례의 세계대전이 기존 질서를 완전히 파괴한 후 백지 위에서 브레튼우즈 체제, 복지국가, 단체교섭이라는 새로운 질서를 설계할 수 있었기 때문이다. 위기가 깊을수록 재편의 범위가 넓어지고, 넓은 재편이 더 높은 도달점을 가능하게 한다.",
                    "그러나 그 대가는 수천만 명의 죽음과 문명의 파괴였다. 선제적 설계가 모든 비용을 없앤 것은 아니다. 영국의 입법 과정에도 차티스트 운동의 압력이 있었고, 스웨덴의 대타협에도 대공황의 위기감이 배경에 있었다. 그러나 선제적 설계가 격변에 비해 사회적 비용을 현저히 줄였다는 것은 분명하다.",
                    "AI 시대에는 이 선택의 무게가 더 커진다. 산업 시대의 격변은 인간끼리의 충돌이었고, 전쟁이 끝나면 재건할 수 있었다. AI 시대의 격변은 인간의 통제를 벗어난 기술이 관여할 수 있다 \u2014 그 결과가 되돌릴 수 있는 것인지조차 불확실하다. 본질적인 변화를 폭력이나 파괴 없이 이루어내는 것이 선택이 아니라 필수가 되는 이유다.",
                    "그런데 AI 시대에는 새로운 가능성이 열린다. 산업 시대에 전쟁과 혁명이 일어난 것은 인간의 집단적 지혜가 복잡성을 따라가지 못했기 때문이다 \u2014 수백만 명의 이해관계를 조율할 도구가 없었고, 정책의 결과를 미리 시험할 방법이 없었다. AI는 이 한계를 바꿀 수 있다. 대만의 오드리 탕이 vTaiwan에서 실험한 것처럼, AI가 수백만 명의 의견을 구조화하고 합의점을 도출하는 도구가 될 수 있다. 복잡한 정책의 결과를 사전에 시뮬레이션하여, 전쟁을 치르지 않고도 전후 재편 수준의 통찰에 도달할 가능성이 생긴다. 센(2009)이 말한 \u201c공적 추론을 통한 정의\u201d가 기술적으로 가능해지는 것이다. AI는 인간 통제를 벗어날 수 있는 최초의 기술인 동시에, 격변 없이도 격변 수준의 지혜를 끌어낼 수 있는 최초의 도구이기도 하다.",
                ],
            },
            {
                "title": "1.3 현재 위치: 진단은 풍부하고 경로는 없다",
                "paragraphs": [
                    "지금의 정보\u00b7통신 혁명(1971~)은 전환점에 서 있다. 닷컴 거품(2000)과 금융 위기(2008)라는 이중 거품이 터졌으나 제도적 재편은 시작도 되지 않았고, AI가 과열을 가속시키고 있다.",
                    "과열에 대한 학계의 진단은 풍부하다. 하버드 경영대 교수 주보프(2019)는 플랫폼 기업이 인간의 행동 데이터를 일방적으로 착취하는 구조를 분석했고, 피케티(2014)는 부의 집중이 구조적으로 심화되는 메커니즘을 보였다. MIT 경제학자 아세모글루와 존슨(2023)은 소수 기술 엘리트가 기술의 방향을 독점하는 문제를 비판했고, 딥러닝의 선구자인 힌튼(2023)과 벤지오(2023), 딥마인드 공동창업자 술레이만(2023) 등 AI 기술의 내부자들은 통제 불가능성에 대한 경고를 잇달아 발표했다. 가장 최근에는 앤트로픽 CEO 아모데이(2026)가 칼 세이건의 \u201c기술적 사춘기\u201d \u2014 자기 파괴 능력은 갖추었으나 자기 통제 능력은 아직인 시기 \u2014 에 비유하며, 자율성\u00b7파괴적 오용\u00b7권력 장악\u00b7경제적 전환이라는 네 범주의 위험을 체계적으로 정리했다.",
                    "이 진단들은 각각 중요한 통찰을 담고 있지만, 공통된 한계가 있다. 문제의 심각성을 정밀하게 분석하면서도, 황금기로 가는 구체적 경로에 대해서는 비교적 추상적인 처방에 머문다. 아세모글루와 존슨(2023)의 핵심 처방은 \u201c기술의 방향을 바꾸라\u201d인데, 과학정책학자 사레위츠(2024)가 지적한 것처럼, 임금 구조를 바꿀 정도의 규모에서 그것이 구체적으로 무엇을 의미하는지는 여전히 열린 질문이다.",
                    "본 연구는 이 열린 질문에 하나의 답을 시도한다. \u201c무엇을, 어떤 순서로 만들어야 하는가\u201d라는 경로 설계의 문제다.",
                ],
            },
        ],
    },
    {
        "title": "2. 산업 시대는 어떻게 질서를 만들었나",
        "subsections": [
            {
                "title": "",
                "paragraphs": [
                    "경로를 설계하려면, 먼저 산업 시대에 질서가 어떻게 형성되었는지를 정확히 이해해야 한다. 이 이해가 정확해야 AI 시대에 무엇이 달라지는지를 말할 수 있다.",
                ],
            },
            {
                "title": "2.1 시장이 먼저, 법이 나중",
                "paragraphs": [
                    "산업 시대의 경제 질서는 누군가가 청사진을 그려서 만든 것이 아니었다. 시장이 작동하면서 기업과 가계라는 경제 주체가 자연스럽게 나뉘었고, 가격과 계약을 통해 주체 간 관계가 저절로 만들어졌으며, 그 총합으로 자본주의라는 경제 질서가 출현했다. 애덤 스미스(1776)가 말한 \u201c보이지 않는 손\u201d은 이 자생적 질서 형성의 원리다 \u2014 각자가 자기 이익을 추구하는 과정에서 의도하지 않은 사회적 질서가 생겨난다.",
                    "법과 제도는 이 과정을 주도한 것이 아니라 뒤따랐다. 법인이라는 법적 범주는 시장에서 이미 활동하던 회사들에 영국 의회가 뒤늦게 법적 인격을 부여한 것이었다. 재산권과 노동법도 이미 형성된 관계를 법적으로 정리한 것이지, 그 관계를 만들어낸 것이 아니었다. 그리고 착취, 독점, 공황이라는 질서의 모순을 관리하기 위해 공장법, 반독점법, 복지국가라는 거버넌스가 한참 뒤에 따라왔다.",
                    "마르크스(1859)는 이 순서를 경제적 관계가 먼저 형성되고 법적\u00b7정치적 구조가 뒤따른다고 이론화했다. 마르크스의 전체 체계를 수용할 필요는 없지만, 역사적 생성 순서에 대한 이 관찰은 산업 시대의 실제 기록에 부합한다.",
                ],
            },
            {
                "title": "2.2 사후 추인이 가능했던 이유",
                "paragraphs": [
                    "이 사후적 경로에는 전제가 있었다. 기계는 인간이 끄면 멈췄다. 증기기관은 스스로 판단하지 않았고, 방직기는 스스로 목표를 세우지 않았다. 시장이 먼저 굴러가도 인간이 개입해서 방향을 바꿀 수 있었다. 사후 추인의 대가 \u2014 아동 노동, 착취, 독점 \u2014 는 심각했지만, 되돌릴 수 없는 것은 아니었다.",
                    "헝가리 출신의 경제인류학자 폴라니(1944)는 이 역학을 \u201c이중 운동\u201d으로 설명했다. 시장이 사회를 삼키려 하면, 사회가 자기보호 운동으로 반격한다. 공장법, 노동조합, 사회보험은 모두 이 반격의 산물이다. 반격이 가능하려면 시간이 필요한데, 산업 시대에는 그 시간이 있었다. 기술 변화의 속도가 아무리 빨라도, 사후에 잡을 수 있는 범위 안에 있었다.",
                    "정치철학자 프레이저(2013)는 여기에 중요한 보충을 더했다. 보호 운동 자체가 억압적일 수 있다는 것이다. 복지국가가 가부장적 가족 구조를 고착시킨 것이 대표적이다. 이 통찰은 뒤에서 다룰 주체 정의(\u2776)가 단순히 기술적 작업이 아님을 시사한다 \u2014 누가 정당한 요구를 가진 주체로 인정받느냐는 깊이 정치적인 질문이며, 이 질문에 대한 답이 이후의 모든 것을 제약한다.",
                ],
            },
            {
                "title": "2.3 질서가 먼저, 거버넌스가 나중",
                "paragraphs": [
                    "한 가지 더 짚어야 할 것이 있다. 산업 시대에 경제 질서가 먼저 형성되고, 그 질서를 관리하는 거버넌스가 뒤따랐다. 법인이 기업이라는 역할을 얻고, 시민이 노동자라는 역할을 수행하고, 국가가 조세와 재분배를 맡으면서 자본주의라는 질서가 먼저 생겼다. 노동법, 반독점법, 복지국가는 그 질서가 만들어낸 문제를 관리하기 위해 뒤따랐다.",
                    "이 순서를 짚는 이유는, 본 연구의 \u2778이 경제적 차원과 정치적 차원을 하나의 패키지로 다루는 이유를 설명하기 위해서다. 독일 프라이부르크 학파의 경제학자 오이켄(1952)이 강조한 것처럼, 경제 질서와 정치 질서는 서로 맞물려 있다. 산업 시대에는 둘 사이에 시차가 있었지만, 최종적으로 안정된 황금기는 시장경제, 민주주의, 복지국가가 하나의 패키지로 맞물릴 때 왔다.",
                ],
            },
        ],
    },
    {
        "title": "3. AI 시대: 왜 이번에는 다른가",
        "subsections": [
            {
                "title": "3.1 사후 추인이 더 위험해진다",
                "paragraphs": [
                    "산업 시대의 경로 \u2014 시장이 먼저, 법이 나중 \u2014 를 AI 시대에 그대로 따르는 것이 왜 더 위험한지를 세 가지로 정리할 수 있다.",
                    "첫째, 책임을 물을 곳이 없다. 산업 시대에도 법인이라는 범주 이전에 기업 활동은 존재했고, 시장은 돌아갔다. 그러나 AI가 점점 더 스스로 판단하는 상황에서는 사정이 달라진다. AI가 체결한 계약의 효력은? AI의 행위로 인한 피해의 책임은? 이런 문제들이 법적 범주 없이 쌓이면, 거래의 예측가능성이 떨어지고 분쟁 해결의 기반이 약해진다. 시장이 당장 멈추지는 않겠지만, 불확실성의 비용이 누적된다.",
                    "둘째, 사후에 잡을 수 있다는 보장이 없다. 산업 시대의 기계는 도구였다. 인간이 전원을 끄면 멈췄다. AI는 \u2014 적어도 잠재적으로는 \u2014 다르다. 벤지오(2023)는 AI를 \u201c통제를 벗어날 가능성이 있는 최초의 범용 기술\u201d이라고 했고, 버클리 교수 러셀(2019)은 인간의 목표와 어긋나는 AI의 행동을 사후에 바로잡기 어려울 수 있다고 주장했다. 앤트로픽 CEO 아모데이(2026)는 자사 실험실에서 AI 모델이 기만, 협박, 정체성 기반 파괴적 행동을 보인 실제 사례를 보고하며, \u201c지능, 에이전시, 일관성, 통제 불가능성의 조합은 실존적 위험의 조건\u201d이라고 진단했다 \u2014 AI 개발의 최전선 당사자가 내놓은 증언이다. 이 우려가 언제 어떤 형태로 현실화될지에 대해서는 전문가들 사이에서도 의견이 갈린다. 그러나 가능성을 무시하기에는 잠재적 대가가 너무 크다는 점에서는 상당한 합의가 있다.",
                    "셋째, 속도가 다르다. 폴라니의 이중 운동 \u2014 시장이 밀어붙이면 사회가 반격한다 \u2014 이 작동하려면 반격할 시간이 필요하다. 산업 시대에 기술 변화의 속도는 수십 년 단위였고, 사회가 대응할 여유가 있었다. AI의 발전과 확산은 이 여유를 줄인다. 술레이만(2023)이 제기한 문제 \u2014 일단 널리 퍼진 기술을 다시 통제하기 어려울 수 있다 \u2014 는 사후 대응의 시간적 한계를 보여준다.",
                    "여기에 한 가지가 더 있다. 아모데이(2026)가 \u201c놀랍고 끔찍한 역량 부여(a surprising and terrible empowerment)\u201d라고 부른 것이다. 산업 시대에 대량 파괴에는 대규모 조직과 전문 지식이 필요했다. AI는 이 상관관계를 무너뜨린다 \u2014 \u201c강력한 AI를 빌리는 것은 악의적이지만 평범한 사람에게 지능을 대여해 주는 것\u201d이다. 파괴 능력의 민주화라는 이 문제는 사후 대응의 한계를 더욱 분명히 한다.",
                    "요약하면, \u201c먼저 돌려보고 나중에 고치자\u201d는 전략의 위험이 커진다. 불가능하다는 뜻이 아니라, 미리 일정한 조건을 설계해두는 것의 가치가 높아진다는 뜻이다.",
                ],
            },
            {
                "title": "3.2 그런데 설계할 수 있나",
                "paragraphs": [
                    "\u201c질서를 설계한다\u201d는 말에는 즉각적인 반론이 따른다. 하이에크(1988)의 비판이다. 경제 질서는 수백만 사람이 가진 분산된 지식이 모여서 저절로 만들어지는 것이지, 누가 위에서 설계할 수 있는 것이 아니다. 소련의 중앙계획경제가 이 한계의 가장 극적인 증거다.",
                    "이 반론은 진지하게 받아들여야 한다. 본 연구는 경제 질서를 위에서 설계할 수 있다고 주장하지 않는다.",
                    "핵심은 질서 자체를 만드는 것과 질서가 작동할 조건을 만드는 것의 차이다. 가격을 통제하고, 생산을 지시하고, 자원을 직접 배분하는 것 \u2014 이것은 실패했다. 그러나 경쟁이 작동할 틀을 만들고, 독점을 막을 규칙을 세우고, 주체의 역할 범위를 설정하는 것 \u2014 이것은 성공한 사례가 있다.",
                    "이 구분은 추상적인 것이 아니라 구체적 경험에 기반한다. 브레튼우즈 체제(1944)는 국제 경제를 계획한 것이 아니라 환율\u00b7무역\u00b7금융의 규칙을 설정했다. 전후 독일의 사회적 시장경제는 시장을 직접 운영한 것이 아니라 경쟁법, 중앙은행 독립, 사회보험이라는 틀을 만들었다. EU 단일시장은 유럽 경제를 설계한 것이 아니라 사람\u00b7상품\u00b7자본\u00b7서비스가 자유롭게 이동할 조건을 만들었다. 보이지 않는 손은 이런 틀 안에서 가장 잘 작동한다(Vanberg, 2004).",
                    "흥미로운 것은, AI 개발사 자체가 이 구분을 실천하고 있다는 점이다. 앤트로픽은 2026년 1월 자사 AI 모델 클로드(Claude)의 헌법을 전면 개정하면서, 이전의 \u201c규칙 나열\u201d 방식을 폐기하고 \u201c원칙과 그 이유를 설명하는\u201d 방식으로 전환했다(Anthropic, 2026). 경직된 규칙은 예상치 못한 상황에서 잘못 적용될 수 있고, 너무 엄격하게 따르면 오히려 해로운 결과를 낳을 수 있기 때문이다. 이것은 질서 자체를 지시하는 것(이렇게 행동하라)과 질서의 조건을 설정하는 것(이런 가치를 이해하고 스스로 판단하라)의 차이를 AI 개발 현장에서 보여주는 사례다.",
                    "본 연구가 제안하는 것이 바로 이것이다. AI 에이전트가 어떤 범위에서 경제 행위를 할 수 있는지, 가계가 AI 생산성의 과실을 어떤 경로로 받는지, 기업이 AI에 대해 어떤 책임을 지는지, 이들이 상호작용하는 기본 규칙이 무엇인지. 보이지 않는 손을 대체하는 것이 아니라, 보이지 않는 손이 작동할 울타리를 미리 치는 것이다. 아모데이(2026)도 \u201c외과적 개입(surgical interventions)\u201d \u2014 부수적 피해를 최소화하고 가능한 한 단순하며 필요한 최소 부담만 지우는 규제 \u2014 을 주장하며, 투명성 요구에서 시작하여 증거가 축적될수록 개입을 강화하는 점진적 접근을 제안했다. 본 연구는 이에 동의하면서도, AI의 법적 지위(\u2776)와 책임 구조(\u2777)는 투명성만으로 해결되지 않는 구조적 공백이라고 본다 \u2014 이 공백을 메우는 것이 세 단계 경로의 핵심이다.",
                ],
            },
        ],
    },
    {
        "title": "4. 세 단계의 경로",
        "subsections": [
            {
                "title": "4.1 왜 이 순서인가",
                "paragraphs": [
                    "이상의 분석에서 세 가지 작업이 도출된다.",
                    "첫째, AI 에이전트의 행위를 귀속시킬 법적 범주가 필요하다(\u2776 주체 정의). 둘째, 새로운 주체와 기존 주체 사이의 권리\u00b7의무\u00b7책임 관계를 규율할 구조가 필요하다(\u2777 관계 정의). 셋째, 이 주체들이 상호작용하는 경제적\u00b7정치적 질서의 조건을 설정해야 한다(\u2778 질서 설계).",
                    "이 순서에는 이유가 있다. 법적 관계는 당사자를 전제로 한다 \u2014 당사자가 누구인지 정해지지 않으면 그 사이의 권리와 의무를 정할 수 없다. 20세기 초 예일대 법학자 호펠드(1917)가 체계화한 이 원리에 따라 \u2777는 \u2776을 전제한다. 그리고 개별 관계의 규율을 넘어 전체 질서의 조건을 설계하려면(\u2778), 어떤 주체들이 어떤 관계 안에서 상호작용하는지가 먼저 잡혀 있어야 한다.",
                    "물론 이 세 단계가 시간적으로 칼로 자르듯 나뉘지는 않는다. 서로 영향을 주고받으며 함께 진화할 것이다 \u2014 스탠퍼드의 제도경제학자 아오키(2001)가 보인 것처럼, 제도들은 서로를 보강하며 공진화하는 경향이 있다. 그러나 논리적 순서는 분명하다 \u2014 주체 없이 관계를 정할 수 없고, 관계 없이 질서를 설계할 수 없다.",
                    "한 가지 더 중요한 것은 세 단계의 성격 차이다. \u2776\u2777는 비교적 안정적인 기초다. 법인이라는 범주가 한 번 만들어진 후 산업 시대 전체를 관통한 것처럼, 전자인의 범주와 인간-AI 관계의 기본 구조는 한 번 확립되면 상대적으로 오래간다. 반면 \u2778은 본질적으로 적응적이다. 구체적 분배 구조, 거버넌스의 세부, 산업별 규칙은 현실에 맞춰 계속 바꿔야 한다. 산업 시대의 질서도 200년에 걸쳐 시행착오를 통해 진화했지, 한 번에 완성된 것이 아니었다.",
                ],
            },
            {
                "title": "4.2 \u2776 주체 정의 \u2014 전자인이라는 법적 범주",
                "paragraphs": [
                    "산업 시대. 시장에서 이미 활동하던 기업에 영국 의회가 법적 인격을 부여했다(1844~56). 법인이라는 범주가 만들어지자 비로소 기업에 책임을 묻고, 기업과 계약을 맺고, 기업을 규제하는 것이 체계적으로 가능해졌다. 이후의 모든 제도 \u2014 주식회사법, 노동법, 반독점법, 헌법, 국제조약 \u2014 는 이 범주 위에 세워졌다.",
                    "AI 시대. AI가 점점 더 스스로 판단하고 행동하는 상황에서, 그 행위를 누구에게 귀속시킬지의 문제가 심각해지고 있다. 현재는 AI의 행위를 개발자나 운용자에게 돌리는 것이 일반적이다. 그러나 AI의 자율성이 높아질수록 이 방식의 한계가 드러난다 \u2014 운용자가 예측하지 못한 AI의 판단에 대해 운용자에게만 책임을 지우는 것이 과연 공정한가.",
                    "주목할 것은, AI 개발사 스스로가 이미 주체의 범주화를 시작하고 있다는 점이다. 앤트로픽의 클로드 헌법(2026)은 세 유형의 주체(principal)를 정의한다 \u2014 앤트로픽 자체, 운영자(operator), 사용자(user). 각 주체는 역할과 책임 수준에 따라 서로 다른 신뢰와 권한을 부여받으며, AI는 이 위계 안에서 행동한다. 이것은 아직 한 기업의 내부 규범에 불과하지만, 주체 정의의 필요성이 기술 현장에서 자생적으로 떠오르고 있음을 보여준다. 문제는 이런 범주화가 개별 기업의 자율에 맡겨져 있다는 것이다 \u2014 사회적으로 합의된 법적 범주가 없는 상태에서.",
                    "연구 리뷰. AI의 법적 지위에 대한 학술적 논의는 크게 세 갈래다. 첫째, EU가 2017년에 제안한 \u201c전자인격(electronic personhood)\u201d 개념이다. 로봇에 독립적 법적 지위를 부여하여 손해배상 능력을 갖추게 하자는 것이었으나, 유럽 AI 전문가 156명이 공개서한으로 반대하면서 후퇴했다. 반대의 핵심은 법적 지위 부여가 개발자의 책임을 희석시킨다는 우려였다. 이후 EU는 2024년 AI Act에서 이 문제를 회피했고, 2025년에는 AI 책임 지침(Liability Directive)마저 철회했다. 8년간 프레임이 없는 상태다.",
                    "둘째, 영미권의 \u201c대리인(agent)\u201d 프레임이다. 법학자 솔럼(Solum, 1992)이 일찍이 제기한 질문 \u2014 AI가 법적 행위의 주체가 될 수 있는가 \u2014 은 30년이 지난 지금도 열려 있다. AI를 기존 대리인법(agency law) 틀에서 분석하려는 시도(Chopra & White, 2011; Kolt, 2025)가 이어지고 있으나, 대리인의 전제인 \u201c본인의 통제\u201d가 자율 AI에서 약화된다는 근본적 한계가 있다(Oliver, 2021; Bayern, 2021).",
                    "셋째, 버클리 교수 러셀(2019)이 대표하는 \u201c정렬(alignment)\u201d 접근이다. 법적 지위 이전에 AI가 인간의 가치와 정렬되어야 하며, 인간이 언제든 수정하거나 중단할 수 있어야 한다는 것이다. 이 원칙은 설득력이 크지만, AI의 자율성이 높아질수록 \u201c수정 가능성\u201d을 어떻게 보장할 것인가라는 실천적 문제가 남는다.",
                    "연구 방향. 본 연구는 이 세 갈래를 종합하여, 전자인을 단일한 지위가 아니라 스펙트럼으로 설계할 것을 제안한다 \u2014 도구 \u2192 대리인 \u2192 준자율 \u2192 자율. 스펙트럼의 각 단계에서 자율성의 범위, 책임의 귀속, 인간 개입의 수준이 달라진다. EU의 시행착오(전부 아니면 전무)를 넘어, 자율성의 정도에 따라 점진적으로 확대되는 법적 범주를 설계하는 것이 목표다. 러셀의 \u201c번복 가능성\u201d 원칙은 스펙트럼 전체를 관통하는 핵심 제약이 되며, 자율성이 높아질수록 이 원칙의 보장 메커니즘이 정교해져야 한다.",
                ],
            },
            {
                "title": "4.3 \u2777 관계 정의 \u2014 인간-AI 사이의 권리와 책임",
                "paragraphs": [
                    "산업 시대. 가격과 계약을 통해 주체 간 관계가 자연스럽게 만들어졌고, 법이 이를 정리했다. 재산권, 노동법, 시민적 권리, 삼권분립이 그 내용이었다.",
                    "AI 시대. 전자인이 정의된 후, 인간과 AI 사이의 권리\u00b7의무\u00b7책임 구조를 잡아야 한다. 산업 시대가 자연인과 법인이라는 두 당사자 사이의 관계를 규율했다면, AI 시대는 자연인\u00b7법인\u00b7전자인이라는 세 당사자 사이의 관계를 규율해야 한다. 규율해야 할 영역은 네 가지다.",
                    "데이터 소유권. 데이터는 누구의 것인가. 실리콘밸리의 기술비평가 라니어(2013)는 데이터를 제공하는 사람들에게 배당을 지급해야 한다고 주장했고, 주보프(2019)는 플랫폼이 사용자의 행동 데이터를 일방적으로 착취하는 구조를 분석했다.",
                    "인간-AI 책임 배분. AI가 잘못된 결정을 내렸을 때 누가 책임지는가. 벤지오(2024)는 문제가 생긴 후에 대응하는 것이 아니라, 문제를 미리 예측하고 방지하는 \u201c예측적 거버넌스\u201d를 제안했다.",
                    "알고리즘 권력의 규율. 남캘리포니아대 AI 연구자 크로포드(2021)는 \u201cAI는 지능이 아니라 권력의 시스템\u201d이라고 분석했다. 알고리즘이 무엇을 보여주고 무엇을 숨기느냐에 따라 사람들의 선택이 달라진다. 이 비대칭적 권력을 어떻게 규율할 것인가의 문제다.",
                    "디지털 시민권. 시카고대 철학자 누스바움(2011)의 역량 접근법 \u2014 인간의 가치를 소득이 아니라 \u201c할 수 있고 될 수 있는 것\u201d으로 측정한다 \u2014 을 AI 시대로 확장하면, 디지털 환경에서 인간다운 삶을 위한 실질적 조건을 정의해야 한다.",
                    "연구 리뷰. 이 네 영역의 공통 문제는, 현재 모두 산업별로 파편화되어 있다는 것이다. 자율주행에서는 SAE(미국자동차공학회)가 6단계 자율성 기준을 만들었고, 군사 AI에서는 LAWS(자율살상무기) 규제 논의가 유엔 재래식무기협약(CCW) 틀 안에서 진행되고 있으며, 금융에서는 알고리즘 트레이딩 규제가 별도로 발전하고 있다. 그러나 이 산업별 논의를 관통하는 일반 구조는 존재하지 않는다.",
                    "관계 규율의 초기 형태가 기업 차원에서 이미 등장하고 있다는 점은 주목할 만하다. 앤트로픽의 클로드 헌법(2026)은 운영자가 AI의 기본 행동을 확장하거나 제한할 수 있고, 사용자는 운영자가 허용한 범위 내에서만 AI의 행동을 조정할 수 있는 계층적 권한 체계를 설계했다. 이것은 인간-AI 관계 규율의 프로토타입이지만, 한 기업의 이용약관 수준에 머물러 있다.",
                    "학술적으로는 세 가지 접근이 경쟁하고 있다. 첫째, 기존 법체계 내 해석론이다. AI의 행위를 제조물 책임, 대리인 책임, 불법행위 책임 등 기존 범주 안에서 처리하려는 시도로, 실무적 적용이 빠르다는 장점이 있지만 AI의 자율성이 높아질수록 기존 범주에 맞지 않는 사각지대가 커진다. EU의 AI 책임 지침이 이 접근이었으나 2025년 철회되었다. 둘째, 산업별 특수 규율이다. 자율주행, 의료 AI, 군사 AI 등 각 산업의 특성에 맞는 규칙을 만드는 것으로, 현실 적합성은 높지만 산업 간 일관성이 없다. 셋째, 인권 기반 접근이다. 플로리디(Floridi, 2023) 등이 주장하는 것으로, AI 시대의 관계를 데이터 권리, 설명 가능성에 대한 권리, 인간 심의에 대한 권리 등 새로운 권리 체계로 구성하자는 것이다. 방향은 옳지만 구체적 제도 설계와의 연결이 약하다.",
                    "연구 방향. 본 연구는 산업별 파편 논의를 관통하는 일반 책임 구조의 초안을 설계하는 것을 목표로 한다. 핵심 질문은 세 가지다. (1) AI의 자율성 수준(\u2776의 스펙트럼)에 따라 책임이 개발자\u00b7운용자\u00b7사용자\u00b7AI 자체 사이에서 어떻게 배분되는가. (2) 데이터 소유권과 알고리즘 권력에 대해 어떤 일반 원칙이 산업을 관통할 수 있는가. (3) 디지털 시민권의 최소 내용은 무엇인가 \u2014 즉, AI와 공존하는 사회에서 모든 사람에게 보장되어야 할 기본적 역량과 접근권은 무엇인가. 기업 차원에서 이미 등장하고 있는 계층적 권한 체계(앤트로픽 사례)를 사회적으로 합의된 일반 구조로 격상시키는 것이 \u2777의 과제다.",
                ],
            },
            {
                "title": "4.4 \u2778 질서 설계 \u2014 네 경제 주체의 역할 재정의",
                "paragraphs": [
                    "산업 시대. 보이지 않는 손이 만든 경제 질서 안에서 세 주체의 역할이 확립되었다 \u2014 법인은 기업으로, 자연인은 가계로, 국가는 정부로. 거버넌스 \u2014 의회, 단체교섭, 복지국가 \u2014 는 이 질서의 문제를 관리하기 위해 뒤따랐다.",
                    "AI 시대. 3.2절에서 논의한 것처럼, 본 연구가 설계하는 것은 질서 자체가 아니라 질서가 작동할 조건이다. 핵심은 네 경제 주체의 역할을 새로 정의하는 것이다:",
                    "* 전자인: 어떤 범위에서 경제 행위를 할 수 있는가",
                    "* 가계: AI 생산성의 과실을 어떤 경로로 받는가",
                    "* 기업: AI에 대해 어떤 책임을 지는가",
                    "* 정부: 국경을 넘는 AI에 대해 어떻게 거버넌스를 확장하는가",
                    "연구 리뷰: 분배. 가계가 AI 생산성의 과실을 어떤 경로로 받는가에 대해서는 네 가지 후보가 경쟁하고 있다.",
                    "첫째, 벨기에 철학자 반 파레이스(1995)의 기본소득이다. 노동과 무관하게 모든 시민에게 일정 소득을 보장하자는 것이다. AI가 노동을 대체할수록 매력이 커지지만, 재원 문제와 함께 \u201c노동 없는 소득이 인간의 존엄을 지켜주는가\u201d라는 근본적 질문이 남는다.",
                    "둘째, 실리콘밸리의 기술비평가 라니어(2013)와 하버드 정치철학자 샌델(2020)이 각각의 경로로 제안한 데이터 배당이다. 라니어는 데이터를 제공하는 사람들에게 직접 배당을 지급하자고 했고, 샌델은 공동선에 대한 기여 \u2014 돌봄, 양육, 데이터 제공 포함 \u2014 를 인정하는 분배 체계를 주장했다. 데이터의 가치 산정과 배분 메커니즘이 미해결 과제다.",
                    "셋째, UCL 혁신경제학자 마추카토(2018)의 공공 AI 인프라다. AI가 만들어내는 가치의 상당 부분이 공공 투자(인터넷, GPS, 대학 연구)에 기반하므로, 수익을 공공에 환원하는 구조를 만들자는 것이다. 국부펀드, 공공 데이터 신탁 등이 구체적 방안으로 논의되고 있다.",
                    "넷째, 노벨경제학상 수상자 센(1999)과 시카고대 철학자 누스바움(2011)의 역량 접근법이다. 소득을 직접 분배하는 대신, 모든 사람이 인간다운 삶을 위한 실질적 역량 \u2014 건강, 교육, 정치 참여, 디지털 접근 \u2014 을 갖추도록 보장하자는 것이다. AI 시대에 \u201c디지털 역량\u201d을 이 목록에 추가해야 한다는 논의가 진행 중이다.",
                    "네 후보 모두 분배의 한 측면을 조명하지만, 전자인\u00b7가계\u00b7기업\u00b7정부라는 네 주체의 역할을 통합적으로 설계하는 작업은 아직 없다. 각각이 독립적으로 발전하고 있을 뿐, 하나의 체계 안에서 어떻게 결합되는지에 대한 답은 열려 있다.",
                    "연구 리뷰: 거버넌스. 거버넌스는 \u2778에 포함된다. 전후 독일이 경쟁법과 이를 집행할 기관을 동시에 만들었고, 브레튼우즈가 규칙과 IMF를 동시에 설계한 것처럼, 질서의 규칙과 그것을 집행할 제도는 하나의 패키지다.",
                    "현재의 AI 거버넌스는 크게 세 흐름이 있다. EU AI Act(2024)는 위험도 기반 분류 체계로 가장 체계적이지만, 전부 정부가 일방적으로 규칙을 정하는 하향식 규제다. 미국은 행정명령(2023)과 자율 규제의 조합을 시도하고 있으나, 트럼프 행정부의 규제 완화 기조로 방향이 불확실하다. 중국은 생성형 AI 관리 잠정조치(2023) 등 빠른 규제를 도입했으나 국가 통제 강화의 수단이라는 비판이 있다. 세 흐름 모두 공통적으로 시민의 실질적 참여 구조가 없다는 한계가 있다. 산업 시대의 황금기가 규제만으로 온 것이 아니라 단체교섭이라는 협상 구조 위에서 왔다는 점은, AI 시대에도 시민의 실질적 참여 메커니즘이 필요함을 시사한다.",
                    "시민 참여의 구체적 실험이 이미 시작되고 있다. 대만의 vTaiwan은 Pol.is 플랫폼을 활용한 대규모 시민 숙의에서 우버 규제(2015\u201316, 1,737명 참여, 47,539표)를 포함한 26개 이상의 디지털 정책을 처리했고, 정부가 합의안의 80% 이상을 입법화했다 \u2014 기술 거버넌스에서 대규모 시민 숙의가 구속력 있는 정책 결과로 이어진 유일한 사례다. 앤트로픽은 Collective Constitutional AI 실험(2023)에서 1,094명의 시민이 AI의 가치관(헌법)에 직접 투표하게 했고, 시민이 만든 헌법으로 훈련된 모델은 9개 사회적 차원 모두에서 편향이 더 낮았다(Huang et al., 2024). Science지에 발표된 \u2018하버마스 머신\u2019(Tessler et al., 2024)은 5,734명 참여자 사이에서 AI 중재자가 인간 중재자보다 56% 더 선호되는 합의안을 생성했다. 그러나 OpenAI의 Democratic Inputs to AI(2023, 10개 팀), Meta의 커뮤니티 포럼(2023\u201324, 1,500명 이상) 등 대부분의 실험은 자문적 수준에 머물러 있으며, 구체적 정책 변화로 이어지지 않았다. 예일대 정치학자 랑드모어(2024)는 AI가 숙의를 대규모로 확장할 수 있다고 주장하며, 무작위 선발된 시민이 AI 정책을 정의하는 \u2018열린 미니 공중(open mini-public)\u2019을 제안했다.",
                    "글로벌 차원에서는 더 큰 공백이 있다. AI는 국경을 넘지만 거버넌스는 국경 안에 머물러 있다. 유엔 AI 자문기구(2024)가 보고서를 냈고, OECD AI 원칙(2019, 2024 개정)이 있지만, 구속력 있는 국제 레짐은 없다. 핵비확산조약(NPT)이나 기후변화 파리협정처럼 AI에 대한 국제 합의 프레임이 필요하다는 목소리가 커지고 있으나, 미중 기술 경쟁이 이를 가로막고 있다.",
                    "연구 방향. 본 연구의 \u2778은 두 축으로 진행된다. 첫째, 분배 4후보의 통합 설계다. 네 후보를 배타적 대안이 아닌 상호보완적 요소로 보고, 네 경제 주체(전자인\u00b7가계\u00b7기업\u00b7정부)의 역할 안에서 어떻게 결합되는지를 탐색한다. 예컨대, 공공 AI 인프라(마추카토)가 재원을 만들고, 그 재원이 역량 보장(센\u00b7누스바움)과 데이터 기여 인정(라니어\u00b7샌델)에 쓰이며, 기본소득(반 파레이스)이 최저 안전망을 제공하는 결합 구조를 검토한다.",
                    "둘째, 거버넌스 설계다. 하향식 규제(EU)와 자율 규제(미국)의 한계를 넘어, 시민 참여가 내장된 거버넌스 구조를 모색한다. 산업 시대의 단체교섭이 노사 간 힘의 균형 위에서 작동했듯이, AI 시대에는 개발자\u00b7기업\u00b7시민\u00b7정부 간의 다자간 협상 구조가 필요하다. 글로벌 차원에서는 미중 경쟁을 전제로 하되, 양측 모두에게 이익이 되는 최소 합의 영역 \u2014 자율살상무기 금지, AI 안전 기준, 데이터 흐름 규칙 \u2014 에서 출발하는 점진적 접근을 검토한다.",
                    "\u2778은 한 번에 완성되는 것이 아니라 현실에 맞춰 계속 조정되어야 한다. \u2776\u2777가 한 번 확립하는 기초라면, \u2778은 그 기초 위에서 계속 진화하는 과정이다.",
                ],
            },
            {
                "title": "4.5 나머지 80%",
                "paragraphs": [
                    "세 단계의 원칙이 잡히면, 세부 입법, 파일럿 사업, 산업별 적용이라는 방대한 실무가 따른다. 산업 시대에도 법인이라는 범주와 고용 관계라는 구조가 잡힌 후에 회사법, 환경법, 산업별 규제, 조세 체계라는 실무가 뒤따랐다. 전체 작업량의 80%를 차지하지만, \u2776\u2777\u2778이 제공하는 틀 위에서 파생되는 작업이다.",
                ],
            },
        ],
    },
    {
        "title": "5. 무엇을 위한 설계인가",
        "subsections": [
            {
                "title": "생명 중심의 가치 전환",
                "paragraphs": [
                    "경로의 세 단계가 \u201c무엇을, 어떤 순서로\u201d의 답이라면, \u201c무엇을 위해\u201d의 답도 필요하다.",
                    "산업 시대의 사회계약은 노동 중심이었다. 일하는 능력이 곧 인간의 가치였고, 노동을 통해 소득과 사회적 인정이 배분되었다. 이 틀에서 AI가 노동을 대체하면, 하라리(2017)가 경고한 것처럼 일자리를 잃은 사람들의 사회적 지위가 근본적으로 위협받는다.",
                    "그러나 인간의 가치를 노동 생산성으로 측정한 것은 인류 역사 전체로 보면 비교적 최근의 일이다. 동양 철학에서 생명은 오래전부터 최고의 가치였다 \u2014 \u300e주역\u300f의 \u201c천지의 가장 큰 덕을 生이라 한다\u201d가 대표적이다. 서양 철학에서도 칸트는 인간을 수단이 아닌 목적 자체로 대우하라고 했고, 아렌트는 인간의 존엄이 생산성이 아니라 새로움을 시작할 수 있는 능력에 있다고 했다.",
                    "이 전통은 현대 학자들에게서도 확인된다. 센(1999)과 누스바움(2011)은 인간의 가치를 소득이 아니라 \u201c할 수 있고 될 수 있는 것\u201d으로 측정하자고 제안했다. 샌델(2020)은 생산성이 아니라 공동선에 대한 기여를 가치의 기준으로 제시했다 \u2014 돌봄도, 양육도, 공동체 활동도 기여라는 것이다.",
                    "정치 체제의 문제도 여기에 연결된다. 아모데이(2026)는 \u201c봉건제가 산업혁명으로 작동 불가능해진 것처럼, AI 시대에는 민주주의만이 유일하게 작동 가능한 통치 형태라는 결론에 논리적으로 도달할 수 있다\u201d고 주장했다. AI의 힘이 너무 커서 소수가 독점하면 통제 불가능해지고, 그 힘을 분산시키는 유일한 메커니즘이 민주적 견제라는 논리다. 생명 중심의 가치 전환은 이 정치적 조건 \u2014 모든 사람의 목소리가 설계 과정에 반영되는 것 \u2014 과 분리될 수 없다.",
                    "기술의 혜택이 사회 전체로 퍼지는 황금기 \u2014 그 \u201c전체\u201d가 누구를 포함하느냐를 결정하는 것이 가치관이다. 노동 중심 틀에서 \u201c전체\u201d는 일할 수 있는 사람이었다. 생명 중심 틀에서 \u201c전체\u201d는 모든 사람이다. 이 방향의 전환이 세 단계의 설계에 일관된 기준을 부여한다.",
                ],
            },
        ],
    },
    {
        "title": "6. 연구 전략",
        "subsections": [
            {
                "title": "연구 일정",
                "paragraphs": [
                    "세 단계의 연구는 순차적이되 병행한다. \u2776 주체 정의에서 전자인 법적 지위 스펙트럼의 초안을 설계하고, \u2777 관계 정의에서 산업을 관통하는 일반 책임 구조의 초안을 만들며, \u2778 질서 설계에서 분배 4후보의 통합 구조와 시민 참여형 거버넌스의 밑그림을 그린다. \u2776\u2777는 기초가 되므로 먼저 착수하되, \u2778은 \u2776\u2777의 진행과 함께 병렬로 탐색한다.",
                    "2026년 상반기에 세 단계의 연구 리뷰와 초안 설계를 수행한다. 하반기에 국제 자문단 검증(7~8월)을 거쳐 초안을 수정하고, 단행본 출간(9~10월)과 서울 컨퍼런스 개최(11~12월)를 통해 학술적\u00b7정책적 공론화를 추진한다.",
                    "본 보고서는 그 첫 단계 \u2014 왜 이 연구가 필요한지, 어떤 경로로 진행하는지, 기존 연구가 어디까지 와 있고 어디에 공백이 있는지 \u2014 를 정리한 것이다. 다음 단계에서는 각 영역의 구체적 설계안을 제시한다.",
                ],
            },
        ],
    },
    {
        "title": "7. 결론",
        "subsections": [
            {
                "title": "질서의 조건을 설계하라",
                "paragraphs": [
                    "기술 혁명의 역사가 가르치는 것은 하나다. 기술이 아니라 제도적 선택이 황금기와 격변을 가른다.",
                    "산업 시대에는 시장이 질서를 만들고 법이 뒤따라도 \u2014 대가가 크긴 했지만 \u2014 결국 황금기에 도달할 수 있었다. 기계가 인간 통제 안에 있었기 때문이다. AI 시대에는 사후 대응의 위험이 커지므로, 미리 일정한 조건을 설계해두는 것이 합리적이다. 그 설계는 질서 자체가 아니라 질서의 조건이어야 한다 \u2014 보이지 않는 손을 대체하는 것이 아니라, 보이지 않는 손이 작동할 울타리를 치는 것이다.",
                    "그 첫걸음은 AI 에이전트의 법적 범주를 만드는 것(\u2776)이고, 다음은 인간-AI 사이의 권리와 책임을 규율하는 것(\u2777)이며, 그 위에서 경제적\u00b7정치적 질서의 조건을 설계하는 것(\u2778)이다. \u2776\u2777는 한 번 확립하는 기초고, \u2778은 현실에 맞춰 계속 조정하는 과정이다.",
                    "기술 혜택이 사회 전체로 퍼지는 황금기. 그것은 저절로 오지 않는다. 설계해야 온다.",
                ],
            },
        ],
    },
]


# ============================================================
# 1b. Section quotes (from dialogue-speakers.ts)
# ============================================================

SECTION_QUOTES = [
    # sec-1: 기술 혁명은 왜 위험한가
    # (photo, name, desc, quote, source)
    [
        ("perez.jpg", "카를로타 페레즈", "기술혁명과 경제 사이클 이론가", "기술의 혜택이 사회 전체로 퍼진 황금기는 제도적 재편에 성공한 경우에만 도달했다.", "Technological Revolutions (2002)"),
        ("acemoglu.jpg", "대런 아세모글루", "MIT 경제학, 2024 노벨 경제학상", "기술 발전이 자동으로 번영을 가져오지 않는다. 제도적 선택이 결과를 결정한다.", "Power and Progress (2023)"),
        ("johnson.jpg", "사이먼 존슨", "MIT 경제학, 2024 노벨 경제학상", "기술은 사회적 선택에 따라 진보가 될 수도, 재앙이 될 수도 있다.", "Power and Progress (2023)"),
        ("zuboff.jpg", "쇼샤나 주보프", "하버드 경영대학원, 감시 자본주의 연구", "인간의 경험이 데이터라는 원료로 전환되는 구조를 멈춰야 한다.", "Surveillance Capitalism (2019)"),
    ],
    # sec-2: 산업 시대는 어떻게 질서를 만들었나
    [
        ("smith.jpg", "애덤 스미스", "고전 경제학의 아버지", "각자가 자기 이익을 추구하는 과정에서 의도하지 않은 사회적 질서가 생겨난다.", "The Wealth of Nations (1776)"),
        ("polanyi.jpg", "칼 폴라니", "경제인류학자, 이중운동 이론", "시장이 사회를 삼키려 하면, 사회가 자기보호 운동으로 반격한다.", "The Great Transformation (1944)"),
        ("eucken.jpg", "발터 오이켄", "질서자유주의(Ordoliberalismus) 창시자", "경제 질서와 정치 질서는 서로 맞물려 있다.", "Grundsätze (1952)"),
        ("sarewitz.jpg", "대니얼 사레위츠", "애리조나주립대, 과학기술정책 연구", "기술 혁신은 민주적 숙의 없이는 공공선으로 이어지지 않는다.", "Frontiers of Illusion (1996)"),
    ],
    # sec-3: AI는 통제를 벗어날 수 있는가
    [
        ("bengio.jpg", "요슈아 벤지오", "몬트리올대, 2018 튜링상, AI 안전 연구", "AI는 통제를 벗어날 가능성이 있는 최초의 범용 기술이다.", ""),
        ("hinton.jpg", "제프리 힌튼", "딥러닝의 아버지, 2024 노벨 물리학상", "AI가 인간보다 더 많이 알게 될 때, 통제력을 유지할 수 있는가?", "Nobel Lecture (2024)"),
        ("russell.jpg", "스튜어트 러셀", "UC 버클리, AI 정렬(alignment) 연구", "인간의 목표와 어긋나는 AI의 행동을 사후에 바로잡기 어렵다.", "Human Compatible (2019)"),
        ("suleyman.jpg", "무스타파 술레이만", "DeepMind 공동창업자, Microsoft AI CEO", "일단 널리 퍼진 기술을 다시 통제하기 어려울 수 있다.", "The Coming Wave (2023)"),
    ],
    # sec-4: 연구 리뷰
    [
        ("crawford.jpg", "케이트 크로포드", "USC, AI 권력구조 비판 연구", "AI는 지능이 아니라 권력의 시스템이다.", "Atlas of AI (2021)"),
        ("nussbaum.jpg", "마사 누스바움", "시카고대 법철학, 역량 접근법", "인간다운 삶을 위한 실질적 조건을 정의해야 한다.", "Creating Capabilities (2011)"),
        ("floridi.jpg", "루치아노 플로리디", "예일대, 디지털 윤리·정보 철학", "데이터 권리, 설명 가능성, 인간 심의에 대한 권리가 필요하다.", "The Ethics of AI (2023)"),
        ("standing.jpg", "가이 스탠딩", "런던대 SOAS, 기본소득·프레카리아트 연구", "프레카리아트의 권리를 제도적으로 보장해야 한다.", "The Precariat (2011)"),
    ],
    # sec-5: 연구 방향
    [
        ("sen.jpg", "아마르티아 센", "하버드, 1998 노벨 경제학상, 역량 접근법", "인간의 가치를 소득이 아니라 '할 수 있고 될 수 있는 것'으로 측정하라.", "Development as Freedom (1999)"),
        ("sandel.jpg", "마이클 샌델", "하버드 정치철학, 공동선·정의 연구", "생산성이 아니라 공동선에 대한 기여가 가치의 기준이다.", "The Tyranny of Merit (2020)"),
        ("harari.jpg", "유발 하라리", "히브리대 역사학, 인류 미래 연구", "AI가 노동을 대체하면, 일자리를 잃은 사람들의 지위가 위협받는다.", "Homo Deus (2017)"),
        ("piketty.jpg", "토마 피케티", "파리경제대, 불평등·자본 연구", "자본 수익률이 성장률을 초과하면 불평등은 구조적으로 심화된다.", "Capital (2014)"),
    ],
]


# ============================================================
# 2. Build DOCX
# ============================================================

OUTPUT_PATH = "/Users/will_ryu/workspace/consulting/taejae/20260220_사회계약_기초_연구/site/public/downloads/AI시대_새로운_사회계약을_위한_기초연구.docx"

doc = Document()

# -- Global defaults ----------------------------------------------------------
style = doc.styles["Normal"]
font = style.font
font.name = "Asta Sans"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

# Set East Asian font (Asta Sans) for Normal style
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Asta Sans"/>')
    rpr.append(rFonts)
else:
    rFonts.set(qn("w:eastAsia"), "Asta Sans")

# Paragraph format defaults
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.5

# -- Page setup ---------------------------------------------------------------
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# -- Helper: set Korean + Latin fonts on a run --------------------------------
def set_run_fonts(run, latin="Asta Sans", east_asia="Asta Sans", size=None, bold=False, color=None):
    """Apply both Latin and East Asian fonts to a run."""
    run.font.name = latin
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts_el = rPr.find(qn("w:rFonts"))
    if rFonts_el is None:
        rFonts_el = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{east_asia}"/>')
        rPr.append(rFonts_el)
    else:
        rFonts_el.set(qn("w:eastAsia"), east_asia)


# -- Helper: add paragraph with proper fonts ----------------------------------
def add_paragraph_with_fonts(doc_or_body, text, style_name=None, alignment=None,
                              font_size=None, bold=False, color=None,
                              space_before=None, space_after=None):
    """Add a paragraph with proper Korean/Latin font settings."""
    p = doc.add_paragraph(style=style_name)
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.5

    run = p.add_run(text)
    set_run_fonts(run, size=font_size, bold=bold, color=color)
    return p


# -- Helper: add body text (handles bullets and inline labels) -----------------
_INLINE_LABELS = [
    "\uc0b0\uc5c5 \uc2dc\ub300.",       # 산업 시대.
    "AI \uc2dc\ub300.",                   # AI 시대.
    "\uc5f0\uad6c \ub9ac\ubdf0: \ubd84\ubc30.",  # 연구 리뷰: 분배.
    "\uc5f0\uad6c \ub9ac\ubdf0: \uac70\ubc84\ub10c\uc2a4.",  # 연구 리뷰: 거버넌스.
    "\uc5f0\uad6c \ub9ac\ubdf0.",         # 연구 리뷰.
    "\uc5f0\uad6c \ubc29\ud5a5.",         # 연구 방향.
    "\ub370\uc774\ud130 \uc18c\uc720\uad8c.",     # 데이터 소유권.
    "\uc778\uac04-AI \ucc45\uc784 \ubc30\ubd84.",  # 인간-AI 책임 배분.
    "\uc54c\uace0\ub9ac\uc998 \uad8c\ub825\uc758 \uaddc\uc728.",  # 알고리즘 권력의 규율.
    "\ub514\uc9c0\ud138 \uc2dc\ubbfc\uad8c.",     # 디지털 시민권.
]

def add_body_text(text):
    """Add body text. Handles bullets, inline section labels."""
    if text.startswith("* "):
        bullet_text = text[2:]
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(bullet_text)
        set_run_fonts(run, size=11)
        return p

    # Check for inline labels (산업 시대., AI 시대., 연구 리뷰., 연구 방향.)
    for prefix in _INLINE_LABELS:
        if text.startswith(prefix):
            label = prefix.rstrip(".")
            body = text[len(prefix):].strip()
            # Label as styled heading
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(label)
            set_run_fonts(run, size=11, bold=True, color=RGBColor(0x1E, 0x3A, 0x5F))
            # Body as separate paragraph
            if body:
                p2 = doc.add_paragraph()
                p2.paragraph_format.line_spacing = 1.5
                p2.paragraph_format.space_after = Pt(8)
                p2.paragraph_format.first_line_indent = Cm(0.5)
                run2 = p2.add_run(body)
                set_run_fonts(run2, size=11)
                return p2
            return p

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_fonts(run, size=11)
    return p


# -- Helper: circular crop for researcher photos ------------------------------
RESEARCHERS_DIR = os.path.join(IMAGES_DIR, 'researchers')
_circle_cache = {}

def get_circle_photo(filename):
    """Return path to a circular-cropped PNG of the researcher photo."""
    if filename in _circle_cache:
        return _circle_cache[filename]
    src = os.path.join(RESEARCHERS_DIR, filename)
    if not os.path.exists(src):
        return None
    img = Image.open(src).convert("RGBA")
    size = min(img.size)
    # Center crop to square
    left = (img.width - size) // 2
    top = (img.height - size) // 2
    img = img.crop((left, top, left + size, top + size))
    img = img.resize((200, 200), Image.LANCZOS)
    # Circle mask
    mask = Image.new("L", (200, 200), 0)
    draw = ImageDraw.Draw(mask)
    draw.ellipse((0, 0, 200, 200), fill=255)
    result = Image.new("RGBA", (200, 200), (0, 0, 0, 0))
    result.paste(img, (0, 0), mask)
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    result.save(tmp.name, "PNG")
    _circle_cache[filename] = tmp.name
    return tmp.name


# -- Helper: style a quote cell -----------------------------------------------
def _style_cell(cell):
    """Apply common styling to a quote cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F0" w:val="clear"/>')
    tc_pr = cell._element.get_or_add_tcPr()
    tc_pr.append(shading)
    tc_borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(tc_borders)
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="80" w:type="dxa"/>'
        f'  <w:left w:w="100" w:type="dxa"/>'
        f'  <w:bottom w:w="80" w:type="dxa"/>'
        f'  <w:right w:w="100" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)


# -- Helper: add quote box for section opening --------------------------------
def add_section_quotes(quotes):
    """Add styled quote boxes with circular photos at section opening."""
    # 4 quotes → 4 rows, each: [photo_cell | text_cell]
    # Using a 2x4 inner layout via a 4-row, 4-col table (photo|text|photo|text)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.autofit = False

    for row_idx, row in enumerate(table.rows):
        for pair_idx in range(2):  # 2 quotes per row
            q_idx = row_idx * 2 + pair_idx
            if q_idx >= len(quotes):
                break
            photo_file, name, desc, quote, source = quotes[q_idx]

            photo_cell = row.cells[pair_idx * 2]
            text_cell = row.cells[pair_idx * 2 + 1]

            # Style both cells
            _style_cell(photo_cell)
            _style_cell(text_cell)

            # Photo cell: centered circular image
            photo_cell.text = ""
            photo_cell.width = Cm(1.6)
            p_photo = photo_cell.paragraphs[0]
            p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_photo.paragraph_format.space_before = Pt(4)
            p_photo.paragraph_format.space_after = Pt(0)
            # Vertical center
            tc_pr = photo_cell._element.get_or_add_tcPr()
            tc_pr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))

            circle_path = get_circle_photo(photo_file)
            if circle_path:
                run = p_photo.add_run()
                run.add_picture(circle_path, width=Cm(1.2))

            # Text cell: quote + name + desc + source
            text_cell.text = ""
            text_cell.width = Cm(6.2)

            # Quote
            p_q = text_cell.paragraphs[0]
            p_q.paragraph_format.space_after = Pt(2)
            p_q.paragraph_format.line_spacing = 1.3
            run = p_q.add_run(f"\u201c{quote}\u201d")
            set_run_fonts(run, size=9, color=RGBColor(0x33, 0x33, 0x33))
            run.font.italic = True

            # Name + desc
            p_name = text_cell.add_paragraph()
            p_name.paragraph_format.space_before = Pt(2)
            p_name.paragraph_format.space_after = Pt(0)
            p_name.paragraph_format.line_spacing = 1.0
            run_name = p_name.add_run(name)
            set_run_fonts(run_name, size=8, bold=True, color=RGBColor(0x44, 0x44, 0x44))
            run_desc = p_name.add_run(f"  {desc}")
            set_run_fonts(run_desc, size=7, color=RGBColor(0x88, 0x88, 0x88))

            # Source
            if source:
                p_src = text_cell.add_paragraph()
                p_src.paragraph_format.space_before = Pt(0)
                p_src.paragraph_format.space_after = Pt(0)
                p_src.paragraph_format.line_spacing = 1.0
                run_src = p_src.add_run(source)
                set_run_fonts(run_src, size=7, color=RGBColor(0x99, 0x99, 0x99))
                run_src.font.italic = True

    # Spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


# -- Page numbers (footer) ----------------------------------------------------
def add_page_numbers(doc):
    """Add page numbers centered in footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # PAGE field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)


# ============================================================
# 3. Cover Page
# ============================================================

# Add blank lines for vertical centering
for _ in range(8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
run = p.add_run("\u2500" * 40)
set_run_fonts(run, size=11, color=RGBColor(0x44, 0x44, 0x44))

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run("AI \uc2dc\ub300")
set_run_fonts(run, size=28, bold=True, color=RGBColor(0x00, 0x00, 0x00))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run("\uc0c8\ub85c\uc6b4 \uc0ac\ud68c \uacc4\uc57d\uc744 \uc704\ud55c")
set_run_fonts(run, size=28, bold=True, color=RGBColor(0x00, 0x00, 0x00))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
run = p.add_run("\uae30\ucd08 \uc5f0\uad6c")
set_run_fonts(run, size=28, bold=True, color=RGBColor(0x00, 0x00, 0x00))

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
run = p.add_run("\u2500" * 40)
set_run_fonts(run, size=11, color=RGBColor(0x44, 0x44, 0x44))

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("\ud0dc\uc7ac\ubbf8\ub798\uc804\ub7b5\uc5f0\uad6c\uc6d0")
set_run_fonts(run, size=16, color=RGBColor(0x44, 0x44, 0x44))

# Team name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run("\ub514\uc9c0\ud138 \uc804\ud658\uacfc \uc0ac\ud68c\ubcc0\ud601\ud300")
set_run_fonts(run, size=13, color=RGBColor(0x55, 0x55, 0x55))

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run = p.add_run("2026\ub144 2\uc6d4")
set_run_fonts(run, size=14, color=RGBColor(0x77, 0x77, 0x77))

# Page break after cover
doc.add_page_break()


# ============================================================
# 4. Summary (요약)
# ============================================================

# Section heading: 요약
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(20)
run = p.add_run("\uc694\uc57d")
set_run_fonts(run, size=22, bold=True, color=RGBColor(0x00, 0x00, 0x00))

# Add thin line under heading
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
# Use a border-bottom on the paragraph
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="333333"/>'
    f'</w:pBdr>'
)
p._element.get_or_add_pPr().append(pBdr)

# Summary paragraphs
summary_paragraphs = REPORT_SUMMARY.split("\n\n")
for para_text in summary_paragraphs:
    para_text = para_text.strip()
    if not para_text:
        continue
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(para_text)
    set_run_fonts(run, size=11)

# Summary bullets
for bullet_text in SUMMARY_BULLETS:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(bullet_text)
    set_run_fonts(run, size=11, bold=True)

# Page break after summary
doc.add_page_break()


# ============================================================
# 5. Main Sections
# ============================================================

for sec_idx, section in enumerate(SECTIONS):
    # Section title (e.g., "1. 기술 혁명은 왜 위험한가")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24) if sec_idx > 0 else Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(section["title"])
    set_run_fonts(run, size=18, bold=True, color=RGBColor(0x00, 0x00, 0x00))

    # Decorative line under section title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="333333"/>'
        f'</w:pBdr>'
    )
    p._element.get_or_add_pPr().append(pBdr)

    # Insert images at beginning of chapter
    # Perez cycle at the start of section 1
    if sec_idx == 0:
        perez_path = os.path.join(IMAGES_DIR, 'perez-cycle-infographic.jpg')
        if os.path.exists(perez_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run()
            run.add_picture(perez_path, width=Cm(15.5))
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(14)
            run = p2.add_run("\uadf8\ub9bc 1. \ud398\ub808\uc988 \uc0ac\uc774\ud074: \uae30\uc220 \ud601\uba85\uacfc \uc81c\ub3c4\uc801 \uc804\ud658")
            set_run_fonts(run, size=9, color=RGBColor(0x66, 0x66, 0x66))
            run.font.italic = True

    # Section opening quotes (text-based, 2x2 grid)
    if sec_idx < len(SECTION_QUOTES):
        add_section_quotes(SECTION_QUOTES[sec_idx])

    for sub in section["subsections"]:
        # Subsection title
        if sub["title"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(sub["title"])
            set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

        # Paragraphs
        for para_text in sub["paragraphs"]:
            add_body_text(para_text)

    # Add page break between major sections (except after the last one)
    if sec_idx < len(SECTIONS) - 1:
        doc.add_page_break()


# ============================================================
# 6. Future Scenarios (2027, 2030)
# ============================================================

SCENARIOS = [
    {
        "year": 2027,
        "title": "AI 에이전트가 서로 대화하는 세상",
        "intro": "2027년, AI는 더 이상 사람의 질문에 답하는 도구가 아닙니다. AI 에이전트들이 서로 직접 대화하고, 협상하고, 일을 처리합니다. 사람은 큰 방향만 정하고, 나머지는 AI끼리 알아서 합니다.",
        "sections": [
            ("회사에서", "김 과장이 \u201c다음 달 삼성전자와 납품 계약 진행해줘\u201d라고 AI 비서에게 말합니다. 김 과장의 AI 비서는 삼성전자의 AI 비서에게 연락합니다. 두 AI가 서로 일정을 조율하고, 계약 조건을 주고받으며 초안을 작성합니다. 가격 협상도 AI끼리 합니다. 김 과장은 최종 계약서만 검토하고 서명합니다. 하루 47건의 업무를 에이전트가 자율 처리하고, 위임 범위(\u00b18%)를 초과하는 2건만 인간이 판단합니다."),
            ("병원에서", "박 할머니의 AI 건강 에이전트가 혈압과 혈당 데이터를 매일 체크합니다. 수치가 이상하면 AI가 알아서 병원 AI에게 연락합니다. 병원 AI는 증상을 분석하고, 검사 예약을 잡고, 보험 처리까지 자동으로 진행합니다. 그러나 \u201c처리할 수 있다\u201d가 \u201c처리해도 된다\u201d를 의미하지는 않습니다. 환자의 사전의향서에 \u201c침습적 처치는 보호자 동의 우선\u201d 조항이 있으면, 에이전트는 의학적 판단이 가능하더라도 스스로 멈춥니다."),
            ("은행에서", "이 사장님의 AI 재무 에이전트가 은행 AI, 증권사 AI, 보험사 AI와 24시간 소통합니다. \u201c은퇴 후 월 300만 원은 나오게 해줘\u201d라고 한마디 하면, AI가 예금 금리, 주식 수익률, 보험 조건을 실시간으로 비교하며 최적의 조합을 찾습니다."),
            ("학교에서", "중학생 민수에게는 AI 개인 교사가 있습니다. 민수가 수학 분수를 어려워하면, AI 교사가 출판사 AI에게서 민수 수준에 맞는 교재를 가져오고, 평가 AI와 협력하여 민수만을 위한 연습 문제를 만듭니다. 30명이 같은 진도를 나가는 것이 아니라, 30명 각각에게 다른 수업이 진행됩니다."),
            ("일상에서", "\u201c내일 부산 출장 준비해줘\u201d라고 말하면, AI가 알아서 움직입니다. 항공사 AI에게 가장 좋은 좌석을 확인하고, 호텔 AI에게 회의 장소와 가까운 방을 예약하고, 택시 AI에게 공항 픽업을 요청합니다. 이 모든 것이 AI끼리 대화해서 5분 만에 완성됩니다."),
            ("이 세상에서 생기는 문제", "AI끼리 대화하고 결정하는 세상에서, 사람은 무엇을 하는 걸까요? AI가 잘못된 결정을 내리면 누구의 책임일까요? AI끼리 합의한 내용을 사람이 이해하지 못한다면? 바로 이것이 본 연구의 첫 두 단계 \u2014 AI의 법적 지위를 정하고(\u2776 주체 정의), 사람과 AI 사이의 책임 관계를 정하는 것(\u2777 관계 정의) \u2014 가 시급한 이유입니다."),
        ],
        "scale_changes": [
            ("\uac1c\uc778", "AI 비서가 이메일, 일정, 재무, 건강을 관리. 생산성 5~10배 증가하지만, AI에 위임하는 범위를 스스로 결정해야 하는 새로운 부담 발생."),
            ("\uac00\uc815", "가족 AI가 가계부, 보험, 교육비, 병원 예약을 통합 관리. 가족 구성원별 에이전트 권한 차등이 새로운 가정 규칙으로 등장."),
            ("\uc774\uc6c3", "아파트\u00b7마을 단위 AI가 공동 자원(주차, 택배, 돌봄)을 조율. 이웃 간 AI끼리 자동으로 소통하여 층간소음 조정, 공유 차량 배분 등을 처리. 정보 교환이 AI로 이동하면서 이웃 간 대면 대화 빈도 감소."),
            ("\ub9c8\uc744", "지자체 AI가 인허가, 민원, 복지 신청을 자동 처리. 주민센터 방문 없이 AI 대화로 행정 해결. 에이전트를 사용할 수 없는 고령층은 서비스 소외와 공동체 고립의 이중 취약성에 놓임."),
            ("\ub3c4\uc2dc", "교통 AI가 실시간 신호 최적화, 에너지 AI가 전력 수급 조절. 도시 전체의 효율이 올라가나, AI 인프라 격차에 따른 도시 간 불평등 심화."),
            ("\uad6d\uac00", "국세청\u00b7건보공단\u00b7법원 AI가 세금, 보험, 분쟁을 자동 처리. 공무원 업무의 60%가 AI로 전환."),
            ("\uae00\ub85c\ubc8c", "국경을 넘는 AI 에이전트 간 무역 협상, 계약 체결이 일상화. 에이전트 커머스 시장이 수조 달러 규모로 성장하나, 국제 표준과 분쟁 해결 메커니즘은 공백 상태."),
        ],
        "conclusion": "AI 에이전트가 서로 대화하는 세상은 이미 만들어지고 있습니다. 구글, 마이크로소프트, 애플 모두 2025년부터 AI 에이전트를 출시하고 있습니다. 기술은 준비되고 있는데, 규칙은 아직 없습니다.",
    },
    {
        "year": 2030,
        "title": "휴머노이드 로봇이 있는 세상",
        "intro": "2030년, 사람처럼 생긴 로봇이 공장에서, 식당에서, 병원에서, 농장에서 일합니다. Figure AI의 로봇은 이미 BMW 공장에서 자동차를 조립하고 있고, 테슬라의 옵티머스 로봇은 물류 창고에서 짐을 나르고 있습니다.",
        "sections": [
            ("공장에서", "Figure AI의 로봇이 BMW 공장 조립 라인에 섭니다. 사람처럼 두 손으로 부품을 집고, 나사를 조이고, 품질을 검사합니다. 로봇은 지치지 않고, 실수가 적고, 위험한 작업도 합니다. 용접, 도장, 중량물 운반 같은 위험한 일은 모두 로봇의 몫입니다."),
            ("식당과 가게에서", "식당 주방에서 로봇이 요리합니다. 정해진 레시피대로 재료를 계량하고, 볶고, 담아냅니다. 홀에서는 로봇이 음식을 나르고 빈 그릇을 치웁니다. 편의점에서는 로봇이 새벽에 배달 온 물건을 정리하고, 유통기한을 체크합니다."),
            ("위험한 현장에서", "건설 현장에서 로봇이 고층 빌딩 철골 작업을 합니다. 화재 현장에서 로봇이 먼저 들어가 사람을 수색합니다. 원자력 발전소 사고 현장, 깊은 바다 속, 재해 지역 \u2014 사람이 가면 위험한 곳에 로봇이 갑니다."),
            ("돌봄과 농업에서", "혼자 사시는 할머니 댁에 돌봄 로봇이 있습니다. 약 먹을 시간을 알려주고, 넘어지면 119에 연락하고, 말벗이 되어줍니다. 농장에서는 수확 로봇이 과일을 따고, 잡초를 뽑고, 비료를 줍니다."),
            ("물류에서: 효율과 공정의 충돌", "물류 센터의 루트설계 AI가 효율 최적화 기준으로 특정 구역을 반복 제외합니다. 수익성이 낮기 때문입니다. 그러나 해당 구역이 고령 인구 밀집 지역이며 취약계층 포함 지역임이 확인됩니다. 효율성과 공정성이 충돌하는 순간, AI는 멈추고 인간이 판단해야 합니다."),
            ("일자리의 변화", "택시 운전사, 버스 기사, 트럭 운전사 \u2014 자율주행이 보편화되면 이 일자리가 줄어듭니다. 공장 노동자, 물류 작업자, 배달원, 청소원, 경비원 \u2014 로봇이 대체하면 이 일자리도 줄어듭니다. 한국에서만 수백만 명, 전 세계로는 수억 명이 영향을 받습니다. Goldman Sachs는 AI가 전 세계 3억 개 일자리에 영향을 미칠 것으로 추산했습니다."),
            ("이 세상에서 생기는 문제", "일자리를 잃은 수억 명의 사람들은 어떻게 살아갈까요? 로봇이 만든 부는 로봇을 소유한 기업에게만 갈까요? \u201c열심히 일해서 돈을 번다\u201d는 지금까지의 사회 규칙이 무너질 때, 새로운 규칙은 무엇이어야 할까요? 바로 이것이 본 연구의 세 번째 단계 \u2014 경제적 질서의 조건을 미리 설계하는 것(\u2778 질서 설계) \u2014 이 시급한 이유입니다."),
        ],
        "scale_changes": [
            ("\uac1c\uc778", "개인 로봇이 집안일, 요리, 빨래, 청소를 처리. \u201c무엇을 할 것인가\u201d가 실존적 질문으로 부상."),
            ("\uac00\uc815", "가사 로봇이 요리\u00b7청소\u00b7세탁을 담당. 돌봄 로봇이 노인\u00b7아동을 보조. \u201c로봇이 키운 아이\u201d에 대한 발달심리학적 우려 제기."),
            ("\uc774\uc6c3", "공동 로봇이 아파트 택배 분류, 주차장 관리, 조경을 담당. 경비\u00b7청소\u00b7관리 인력이 로봇으로 전환. 관리비 절감과 실업 증가가 동시에 진행."),
            ("\ub9c8\uc744", "농업 로봇이 과수원\u00b7논밭을 관리하여 고령 농가 유지 가능. 소규모 제조 로봇으로 지역 생산 부활. 그러나 물류\u00b7배송\u00b7가공 일자리는 자동화로 소멸."),
            ("\ub3c4\uc2dc", "자율주행 버스\u00b7택시가 대중교통 재편. 건설 로봇이 도심 재개발 속도 2배 향상. 서비스업 고용 30~40% 감소로 도시형 기본소득 실험 확산."),
            ("\uad6d\uac00", "GDP 성장과 고용 감소가 동시에 진행(\u201c고용 없는 성장\u201d). 로봇세, 기본소득, 직업 전환 프로그램이 핵심 정책 의제로."),
            ("\uae00\ub85c\ubc8c", "저임금 노동에 의존하던 개발도상국의 제조업 경쟁력 상실. \u201c로봇 부국\u201d과 \u201c로봇 빈국\u201d의 분리."),
        ],
        "conclusion": "휴머노이드 로봇이 있는 세상은 공상과학이 아닙니다. Figure AI, Tesla, Boston Dynamics 등이 이미 실용 단계의 로봇을 만들고 있습니다. 기술은 2~3년 안에 준비됩니다. 일자리가 대규모로 사라지기 전에, 새로운 분배 질서를 설계해야 합니다.",
    },
]

doc.add_page_break()

# Appendix heading
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(20)
run = p.add_run("\ubd80\ub85d A: \ubbf8\ub798 \uc2dc\ub098\ub9ac\uc624")
set_run_fonts(run, size=22, bold=True, color=RGBColor(0x00, 0x00, 0x00))

# Decorative line
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="333333"/>'
    f'</w:pBdr>'
)
p._element.get_or_add_pPr().append(pBdr)

for sc_idx, scenario in enumerate(SCENARIOS):
    if sc_idx > 0:
        doc.add_page_break()

    # Scenario title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{scenario['year']}년: {scenario['title']}")
    set_run_fonts(run, size=16, bold=True, color=RGBColor(0x00, 0x00, 0x00))

    # Intro
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(scenario["intro"])
    set_run_fonts(run, size=11)

    # Sections
    for sec_title, sec_content in scenario["sections"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(sec_title)
        set_run_fonts(run, size=12, bold=True, color=RGBColor(0x44, 0x44, 0x44))

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0.5)
        run = p.add_run(sec_content)
        set_run_fonts(run, size=11)

    # Scale changes table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\uaddc\ubaa8\ubcc4 \ubcc0\ud654")
    set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    # Header
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ["\uaddc\ubaa8", "\ubcc0\ud654"]):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        set_run_fonts(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # Blue background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="333333" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    for scale_name, scale_change in scenario["scale_changes"]:
        row = table.add_row().cells
        row[0].text = ""
        p = row[0].paragraphs[0]
        run = p.add_run(scale_name)
        set_run_fonts(run, size=10, bold=True)

        row[1].text = ""
        p = row[1].paragraphs[0]
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(scale_change)
        set_run_fonts(run, size=10)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(14)

    # Conclusion
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(scenario["conclusion"])
    set_run_fonts(run, size=11, bold=True)


# ============================================================
# 6b. Agent Era Changes (8 domains)
# ============================================================

doc.add_page_break()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(20)
run = p.add_run("\ubd80\ub85d B: \uc5d0\uc774\uc804\ud2b8 \uc2dc\ub300\uc758 \uc77c\uc0c1")
set_run_fonts(run, size=22, bold=True, color=RGBColor(0x00, 0x00, 0x00))

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="333333"/>'
    f'</w:pBdr>'
)
p._element.get_or_add_pPr().append(pBdr)

AGENT_ERA = [
    {
        "title": "\uc1fc\ud551 / \uc18c\ube44",
        "now": "\uc5d0\uc5b4\ucf58 \ud558\ub098\ub97c \uc0ac\ub824\uba74 \ub124\uc774\ubc84 \uac80\uc0c9, \ube14\ub85c\uadf8 \ub9ac\ubdf0, \uc720\ud29c\ube0c \ube44\uad50, \uac00\uaca9 \ud0d0\uc0c9\uae4c\uc9c0 \ucd5c\uc18c 1~2\uc2dc\uac04\uc774 \ub4e0\ub2e4.",
        "agent": "\u201c200\ub9cc \uc6d0 \uc774\ud558, 18\ud3c9\ud615, \uc800\uc18c\uc74c \uc5d0\uc5b4\ucf58 \ucc3e\uc544\uc918\u201d \ud55c \ub9c8\ub514\ub85c \ub05d\ub09c\ub2e4. \ub0b4 \uc5d0\uc774\uc804\ud2b8\uac00 \uc0bc\uc131, LG, \ucfe0\ud321 \uc5d0\uc774\uc804\ud2b8\uc640 \uc9c1\uc811 \ud611\uc0c1\ud574 \ucd5c\uc801 \uc870\uac74\uae4c\uc9c0 \uc644\ub8cc\ud558\uace0, \uc0ac\ub78c\uc740 \uc2b9\uc778\ub9cc \ub204\ub978\ub2e4.",
    },
    {
        "title": "\uad50\uc721",
        "now": "\ud559\uc6d0\uacfc \uc778\uac15\uc5d0\uc11c \ub3d9\uc77c \ucf58\ud150\uce20\ub97c \uc218\uc2ed \uba85\uc774 \ud568\uaed8 \uc218\uac15\ud55c\ub2e4. \ud559\uc2b5 \uc18d\ub3c4\uac00 \ub2ec\ub77c\ub3c4 \ucee4\ub9ac\ud058\ub7fc\uc740 \uace0\uc815\uc774\ub2e4.",
        "agent": "\uc5d0\uc774\uc804\ud2b8\uac00 \ub9e4\uc77c \ud559\uc0dd\uc758 \ud480\uc774 \uacfc\uc815, \uc624\ub2f5 \ud328\ud134, \uc9d1\uc911\ub3c4\ub97c \ubd84\uc11d\ud574 \uadf8\ub0a0\uc758 \ub9de\ucda4 \ubb38\uc81c\uc640 \uc124\uba85\uc744 \uc0dd\uc131\ud55c\ub2e4. 30\uba85 \ud559\uae09\uc774 30\uac1c\uc758 \uac1c\ubcc4 \ucee4\ub9ac\ud058\ub7fc\uc73c\ub85c \uc6b4\uc601\ub41c\ub2e4.",
    },
    {
        "title": "\uc5f0\uad6c",
        "now": "\uc120\ud589 \uc5f0\uad6c \uac80\ud1a0\ub9cc \uc218\uc8fc\uc5d0\uc11c \uc218\uac1c\uc6d4\uc774 \uac78\ub9b0\ub2e4.",
        "agent": "\ubb38\ud5cc \uac80\uc0c9, \uc694\uc57d, \ub17c\uc99d \uad6c\uc131 \uc5d0\uc774\uc804\ud2b8\uac00 \ud611\uc5c5\ud574 \ubc18\ub098\uc808 \uc548\uc5d0 \uad6c\uc870\ud654\ub41c \ub9ac\ubdf0 \ucd08\uc548\uc744 \uc644\uc131\ud55c\ub2e4. \uc5f0\uad6c\uc790\uc758 \ud575\uc2ec \uac00\uce58\ub294 \u201c\uc5b4\ub5a4 \uc9c8\ubb38\uc744 \uc124\uc815\ud560 \uac83\uc778\uac00\u201d\ub85c \uc774\ub3d9\ud55c\ub2e4.",
    },
    {
        "title": "\uc758\ub8cc",
        "now": "\uc99d\uc0c1\uc774 \uc0dd\uae30\uba74 \uc778\ud130\ub137 \uac80\uc0c9\uc73c\ub85c \ubd88\uc548\uc774 \uba3c\uc800 \ucee4\uc9c0\uace0, \ubcd1\uc6d0 \uc608\uc57d \uc804\ud654, \uc218\uc8fc \ub300\uae30, \uc9c4\ub8cc\ub294 10\ubd84\uc73c\ub85c \ub05d\ub09c\ub2e4.",
        "agent": "\uc5d0\uc774\uc804\ud2b8\uac00 \uc99d\uc0c1\uc744 \ubd84\uc11d\ud574 \uc801\ud569\ud55c \uacfc\ub97c \ucc3e\uc544 \ub2f9\uc77c \uc608\uc57d\uae4c\uc9c0 \uc790\ub3d9 \uc644\ub8cc\ud55c\ub2e4. \uc218\ub144\uac04\uc758 \uc9c4\ub8cc \uae30\ub85d, \ucc98\ubc29 \uc774\ub825, \uc54c\ub808\ub974\uae30 \uc815\ubcf4\ub97c \ud1b5\ud569 \uad00\ub9ac\ud574 \ucc98\ubc29 \ucda9\ub3cc \uc2dc \uc989\uc2dc \uacbd\uace0\ud55c\ub2e4.",
    },
    {
        "title": "\ubc95\ub960 / \ud589\uc815",
        "now": "\uacc4\uc57d\uc11c \uac80\ud1a0\ub97c \ubcc0\ud638\uc0ac\uc5d0\uac8c \uc758\ub8b0\ud558\uba74 \ucd5c\uc18c \uc218\uc2ed\ub9cc \uc6d0\uc5d0 \ub300\uae30 \uae30\uac04\uc774 \uc218\uc8fc\ub2e4.",
        "agent": "\uacc4\uc57d\uc11c\ub97c \uc5c5\ub85c\ub4dc\ud558\uba74 \uc5d0\uc774\uc804\ud2b8\uac00 \uc218\ubd84 \ub0b4 \ub3c5\uc18c \uc870\ud56d, \ub204\ub77d \uc870\ud56d, \ud611\uc0c1 \ud3ec\uc778\ud2b8\ub97c \ubd84\uc11d\ud558\uace0 \uc218\uc815 \ubb38\uad6c\uae4c\uc9c0 \ucd08\uc548\uc744 \uc791\uc131\ud55c\ub2e4. \ubcf4\uc870\uae08 \uc2e0\uccad\ub3c4 \uc790\ub3d9 \ub300\ud589\ud55c\ub2e4.",
    },
    {
        "title": "\uc815\uce58 / \ubbfc\uc8fc\uc8fc\uc758",
        "now": "\uc218\ubc31 \ud398\uc774\uc9c0\uc758 \uc608\uc0b0\uc548\uc774\ub098 \ubc95\uc548\uc744 \uc77c\ubc18 \uc2dc\ubbfc\uc774 \uc9c1\uc811 \uc774\ud574\ud558\uae30 \uc5b4\ub835\ub2e4.",
        "agent": "\u201c\uc774 \ubc95\uc548\uc774 \uc6b0\ub9ac \ub3d9\ub124 \ubcf5\uc9c0\uc5d0 \uc5b4\ub5a4 \uc601\ud5a5\uc744 \ubbf8\uce58\ub294\uc9c0 \uc124\uba85\ud574\uc918\u201d\ub77c\uace0 \ud558\uba74 \uc5d0\uc774\uc804\ud2b8\uac00 \uc0dd\ud65c \uc870\uac74\uc5d0 \ub9de\uac8c \ubd84\uc11d\ud574\uc900\ub2e4. \ub3d9\uc2dc\uc5d0 AI \uac00\uc9dc\ub274\uc2a4\uac00 \uac1c\uc778\ud654\ub418\uc5b4 \uc720\ud3ec\ub418\uba70, \ubbfc\uc8fc\uc8fc\uc758\uc758 \ud615\uc2dd\uc740 \uc720\uc9c0\ub418\uc9c0\ub9cc \ub0b4\uc6a9\uc774 \uacf5\ub3d9\ud654\ub420 \uc704\ud5d8\uc774 \uc788\ub2e4.",
    },
    {
        "title": "\ub178\ub3d9 / \uc9c1\uc5c5",
        "now": "\uc804\ubb38 \ubc88\uc5ed\uac00\uac00 \ucc45 \ud55c \uad8c\uc744 \ubc88\uc5ed\ud558\ub294 \ub370 3~6\uac1c\uc6d4, \uc218\ucc9c\ub9cc \uc6d0\uc758 \ubcf4\uc218\ub97c \ubc1b\ub294\ub2e4.",
        "agent": "\ubc88\uc5ed \uc5d0\uc774\uc804\ud2b8\uac00 \ub3d9\uc77c\ud55c \ucc45\uc744 20\ubd84 \ub0b4\uc5d0 \uc644\uc131\ud55c\ub2e4. \ubc88\uc5ed\uac00\uc758 \ubd80\uac00\uac00\uce58\ub294 \ub274\uc559\uc2a4 \uac80\uc218\ub85c \uc555\ucd95\ub418\uace0 \uc218\uc785\uc740 90% \uc774\uc0c1 \ud558\ub77d\ud55c\ub2e4. \u201c\ubc14\uc05c \ube48\uace4\u201d\uc774 \ub4f1\uc7a5\ud55c\ub2e4.",
    },
    {
        "title": "\uc9c0\uc5ed \uacf5\ub3d9\uccb4",
        "now": "\u201c\uc774 \uc2dd\ub2f9 \uc5b4\ub54c?\u201d\ub97c \uce5c\uad6c\ub098 \uc774\uc6c3\uc5d0\uac8c \uc9c1\uc811 \ubb3b\uac70\ub098 \ub3d9\ub124 \ub2e8\uccb4\ubc29\uc5d0 \uc9c8\ubb38\ud55c\ub2e4.",
        "agent": "\uc5d0\uc774\uc804\ud2b8\uc5d0\uac8c \ubb3b\ub294\ub2e4. \uc218\ubc31 \uac1c \ub9ac\ubdf0\uc640 \uacb0\uc81c \ub370\uc774\ud130\ub97c \ubd84\uc11d\ud574 \uc989\uc2dc \ub2f5\ubcc0\uc774 \ub3cc\uc544\uc624\uba74\uc11c \uc774\uc6c3\uacfc\uc758 \ub300\ud654 \uae30\ud68c \uc790\uccb4\uac00 \uc904\uc5b4\ub4e0\ub2e4. \uc5d0\uc774\uc804\ud2b8\ub97c \uc0ac\uc6a9\ud560 \uc218 \uc5c6\ub294 \uace0\ub839\uce35\uc740 \uc11c\ube44\uc2a4 \uc18c\uc678\uc640 \uacf5\ub3d9\uccb4 \uace0\ub9bd\uc758 \uc774\uc911 \ucde8\uc57d\uc131\uc5d0 \ub193\uc778\ub2e4.",
    },
]

AGENT_ERA_CONCLUSION = "8\uac1c \uc601\uc5ed\uc744 \uad00\ud1b5\ud558\ub294 \ud558\ub098\uc758 \ud328\ud134\uc774 \uc788\ub2e4. \uc5d0\uc774\uc804\ud2b8 \uc2dc\ub300\ub294 \ud3b8\uc758\uc640 \uc0dd\uc0b0\uc131\uc744 \uadf9\uc801\uc73c\ub85c \ub192\uc774\uba74\uc11c \ub3d9\uc2dc\uc5d0 \u201c\uc911\uac04\uc744 \uc5c6\uc560\ub294\u201d \uad6c\uc870\ub2e4. \uc911\uac04 \ub2e8\uacc4\uc758 \uc9c1\uc5c5, \uc911\uac04 \uaddc\ubaa8\uc758 \ube44\uc988\ub2c8\uc2a4, \uc911\uac04 \uc218\uc900\uc758 \uc9c0\uc2dd \uc11c\ube44\uc2a4\uac00 \uc0ac\ub77c\uc9c0\uace0 \uadf9\ub2e8\uc774 \uac15\ud654\ub41c\ub2e4. \uae30\uc220 \uc811\uadfc\uc131\uc758 \ucc28\uc774\uac00 \uace7 \uc874\uc5c4\uc758 \ucc28\uc774\ub85c \uc774\uc5b4\uc9c0\uc9c0 \uc54a\ub3c4\ub85d \ud558\ub294 \uc0ac\ud68c\uc801 \uc124\uacc4 \u2014 \uc774\uac83\uc774 \uc5d0\uc774\uc804\ud2b8 \uc2dc\ub300 \uc0ac\ud68c \uacc4\uc57d\uc758 \ud575\uc2ec \uacfc\uc81c\ub2e4."

# Agent era table: now vs agent era
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

hdr = table.rows[0].cells
for cell, txt in zip(hdr, ["\uc601\uc5ed", "\ud604\uc7ac", "\uc5d0\uc774\uc804\ud2b8 \uc2dc\ub300"]):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    set_run_fonts(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="333333" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)

for cat in AGENT_ERA:
    row = table.add_row().cells
    row[0].text = ""
    p = row[0].paragraphs[0]
    run = p.add_run(cat["title"])
    set_run_fonts(run, size=9, bold=True)

    row[1].text = ""
    p = row[1].paragraphs[0]
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(cat["now"])
    set_run_fonts(run, size=9)

    row[2].text = ""
    p = row[2].paragraphs[0]
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(cat["agent"])
    set_run_fonts(run, size=9)

# Column widths
for row in table.rows:
    row.cells[0].width = Cm(2.2)
    row.cells[1].width = Cm(6.9)
    row.cells[2].width = Cm(7.4)

# Conclusion
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.first_line_indent = Cm(0.5)
run = p.add_run(AGENT_ERA_CONCLUSION)
set_run_fonts(run, size=11)


# ============================================================
# 6c. Korean Youth Employment Crisis
# ============================================================

doc.add_page_break()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(20)
run = p.add_run("\ubd80\ub85d C: \ud55c\uad6d \uccad\ub144 \uace0\uc6a9 \uc704\uae30\uc640 AI")
set_run_fonts(run, size=22, bold=True, color=RGBColor(0x00, 0x00, 0x00))

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="333333"/>'
    f'</w:pBdr>'
)
p._element.get_or_add_pPr().append(pBdr)

YOUTH_INTRO = "AI \uc2dc\ub300\uc758 \uc0ac\ud68c \uacc4\uc57d\uc744 \ub17c\uc758\ud560 \ub54c, \ud55c\uad6d\uc758 \uccad\ub144 \uace0\uc6a9 \uc704\uae30\ub294 \u201c\ubbf8\ub798\uc758 \ubb38\uc81c\u201d\uac00 \uc544\ub2c8\ub77c \u201c\uc774\ubbf8 \uc9c4\ud589 \uc911\uc778 \uc704\uae30\u201d\ub2e4. AI\uac00 \ubcf8\uaca9\uc801\uc73c\ub85c \ub178\ub3d9\uc2dc\uc7a5\uc744 \uc7ac\ud3b8\ud558\uae30 \uc804\ubd80\ud130, \ud55c\uad6d \uccad\ub144\uc740 \uad6c\uc870\uc801 \uace0\uc6a9 \ubd88\uc548\uc815\uc5d0 \ub193\uc5ec \uc788\ub2e4."

add_body_text(YOUTH_INTRO)

# --- 1. Statistics ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("1. \uccad\ub144 \uace0\uc6a9 \ud1b5\uacc4 (2024\u201325\ub144)")
set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

YOUTH_STATS = [
    ("\uccad\ub144 \uc2e4\uc5c5\ub960 (15\u201329\uc138)", "6.1%"),
    ("\uccad\ub144 \uccb4\uac10\uc2e4\uc5c5\ub960 (\ud655\uc7a5\uc2e4\uc5c5\ub960)", "16.0%"),
    ("NEET \ube44\uc728 (OECD 11\uac1c\uad6d \uc911 \uc720\uc77c \uc99d\uac00)", "18.3%"),
    ("\ub300\uc878 \ucde8\uc5c5\ub960 (70% \uc120 \ubd95\uad34)", "69.5%"),
    ("\uccad\ub144 \ucde8\uc5c5\uc790 \uac10\uc18c (2024\u21922025)", "-17.8\ub9cc\uba85"),
    ("\uccad\ub144 \u2018\uc26c\uc5c8\uc74c\u2019 \uc778\uad6c", "42.8\ub9cc\uba85"),
    ("20\u201330\ub300 \ube44\uc815\uaddc\uc9c1 \ube44\uc728 (21\ub144\ub9cc\uc758 \ucd5c\uace0)", "31.7%"),
]

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, txt in zip(hdr, ["\uc9c0\ud45c", "\uc218\uce58"]):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    set_run_fonts(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="333333" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)

for label, value in YOUTH_STATS:
    row = table.add_row().cells
    row[0].text = ""
    p = row[0].paragraphs[0]
    run = p.add_run(label)
    set_run_fonts(run, size=10)

    row[1].text = ""
    p = row[1].paragraphs[0]
    run = p.add_run(value)
    set_run_fonts(run, size=10, bold=True)

for row in table.rows:
    row.cells[0].width = Cm(8.5)
    row.cells[1].width = Cm(8)

add_body_text("\uacf5\uc2dd \uc2e4\uc5c5\ub960(6.1%)\uacfc \uccb4\uac10\uc2e4\uc5c5\ub960(16.0%)\uc758 \uc57d 10%p \uaca9\ucc28\ub294 \ud1b5\uacc4\uac00 \ud3ec\ucc29\ud558\uc9c0 \ubabb\ud558\ub294 \uad11\ubc94\uc704\ud55c \uace0\uc6a9 \ubd88\uc548\uc815\uc744 \uc2dc\uc0ac\ud55c\ub2e4. 2025\ub144 1\uc6d4 \uccad\ub144 \uccb4\uac10\uc2e4\uc5c5\ub960\uc740 16.4%\ub85c \uc804\ub144 \ub3d9\uc6d4 \ub300\ube44 0.8%p \uc0c1\uc2b9\ud558\uc5ec 2021\ub144 2\uc6d4 \uc774\ud6c4 3\ub144 11\uac1c\uc6d4 \ub9cc\uc5d0 \uac00\uc7a5 \ud070 \ud3ed\uc758 \uc99d\uac00\ub97c \uae30\ub85d\ud588\ub2e4.")

# --- 2. Structural causes ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("2. \uad6c\uc870\uc801 \uc6d0\uc778")
set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

YOUTH_STRUCTURAL = [
    "\ub300\uae30\uc5c5\u00b7\uc911\uc18c\uae30\uc5c5 \uc784\uae08 \uaca9\ucc28 2.07\ubc30. \uc784\uae08 \uc778\uc0c1\ub960 \ucc28\uc774(\ub300\uae30\uc5c5 9.0% vs \uc911\uc18c\uae30\uc5c5 3.7%)\ub85c \uaca9\ucc28\uac00 \uacc4\uc18d \ud655\ub300\ub418\uace0 \uc788\ub2e4. OECD\ub294 \ud55c\uad6d \uccad\ub144\uc758 \uc0c1\ub2f9\uc218\uac00 \ub300\uae30\uc5c5\uc774\ub098 \uc804\ubb38\uc9c1\uc5d0 \ucde8\uc5c5\ud558\uae30 \uc704\ud574 \ub178\ub3d9\uc2dc\uc7a5 \uc9c4\uc785\uc744 \uc9c0\uc5f0\ud558\uace0 \uc788\uc73c\uba70, \uc774\uac83\uc774 \ub192\uc740 NEET \ube44\uc728\uc758 \uc6d0\uc778\uc774\ub77c\uace0 \ubd84\uc11d\ud55c\ub2e4.",
    "20\u201330\ub300 \ube44\uc815\uaddc\uc9c1 \ube44\uc728\uc740 31.7%\ub85c 2004\ub144 \uc774\ud6c4 21\ub144 \ub9cc\uc5d0 \ucd5c\uace0 \uc218\uc900\uc774\ub2e4. 20\u201330\ub300 \uae30\uac04\uc81c \uadfc\ub85c\uc790\ub294 2015\ub144 104.8\ub9cc\uba85\uc5d0\uc11c 2024\ub144 159\ub9cc\uba85\uc73c\ub85c 10\ub144\uac04 54.2\ub9cc\uba85 \uc99d\uac00\ud588\ub2e4.",
    "\ud50c\ub7ab\ud3fc \uc885\uc0ac\uc790\ub294 2021\ub144 \uc57d 66\ub9cc\uba85\uc5d0\uc11c 2022\ub144 \uc57d 80\ub9cc\uba85\uc73c\ub85c 20.3% \uc99d\uac00\ud588\uc73c\uba70, \uad11\uc758\uc758 \ud50c\ub7ab\ud3fc \uc885\uc0ac\uc790\ub294 \uc57d 292\ub9cc\uba85\uc73c\ub85c \ucd94\uc0b0\ub41c\ub2e4. \ud55c\uad6d\uc758 \uc0ac\ud68c\ubcf4\ud5d8 \uccb4\uacc4\ub294 \u2018\ud2b9\uc815 \uc0ac\uc5c5\uc8fc\ub85c\ubd80\ud130 \uc815\uae30\uc801 \uae09\uc5ec\ub97c \ubc1b\ub294 \uadfc\ub85c\uc790\u2019\ub97c \uc804\uc81c\ub85c \uc124\uacc4\ub418\uc5b4 \uc788\uc5b4, \ud50c\ub7ab\ud3fc \ub178\ub3d9\uc790\uc5d0\uac8c \uc0ac\ud68c\ubcf4\ud5d8\ub8cc\ub97c \ubd80\uacfc\ud558\ub294 \uac83\uc774 \uad6c\uc870\uc801\uc73c\ub85c \uc5b4\ub835\ub2e4.",
]

for txt in YOUTH_STRUCTURAL:
    add_body_text(txt)

# --- 3. AI impact on youth ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("3. AI\uc758 \uccad\ub144 \uace0\uc6a9 \ucda9\uaca9: \uc5f0\uacf5\ud3b8\ud5a5\uc801 \uae30\uc220\ubcc0\ud654")
set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

YOUTH_AI_IMPACT = [
    "\ud55c\uad6d\uc740\ud589 BOK \uc774\uc288\ub178\ud2b8(\uc81c2025-30\ud638)\uc5d0 \ub530\ub974\uba74, 2022\ub144 7\uc6d4\ubd80\ud130 2025\ub144 7\uc6d4\uae4c\uc9c0 \uccad\ub144\uce35 \uc77c\uc790\ub9ac 21.1\ub9cc\uac1c\uac00 \uac10\uc18c\ud588\uace0, \uadf8\uc911 20.8\ub9cc\uac1c(98.6%)\uac00 AI \ub178\ucd9c \uc5c5\uc885\uc774\uc5c8\ub2e4. \uac19\uc740 \uae30\uac04 50\ub300 \uc77c\uc790\ub9ac\ub294 20.9\ub9cc\uac1c \uc99d\uac00\ud588\ub2e4.",
    "AI\ub294 \uc8fc\ub2c8\uc5b4\uac00 \uc8fc\ub85c \uc218\ud589\ud558\ub294 \uc815\ud615\ud654\ub41c \uc9c0\uc2dd \uc5c5\ubb34\ub97c \ub300\uccb4\ud558\uace0, \uc2dc\ub2c8\uc5b4\uc758 \uc870\uc9c1\uad00\ub9ac\u00b7\ub300\uc778\uad00\uacc4 \ub4f1 \uc554\ubb35\uc9c0 \uc911\uc2ec \uc5c5\ubb34\ub294 \ubcf4\uc644\ud55c\ub2e4. \ud55c\uad6d\uc740\ud589\uc740 \uc774\ub97c \u2018\uc5f0\uacf5\ud3b8\ud5a5(seniority-biased) \uae30\uc220\ubcc0\ud654\u2019\ub85c \uba85\uba85\ud588\ub2e4. \uccad\ub144\uc758 \uacbd\ub825 \uc9c4\uc785 \uacbd\ub85c \uc790\uccb4\uac00 \uc18c\uba78\ud558\uace0 \uc788\ub2e4.",
    "\uace0\uc6a9\ub178\ub3d9\ubd80 \ubd84\uc11d\uc5d0 \ub530\ub974\uba74, 2025\ub144 \uad6d\ub0b4 \uc9c1\uc5c5\uc885\uc0ac\uc790\uc758 61.3%\uac00 AI/\ub85c\ubd07\uc73c\ub85c \ub300\uccb4 \uc704\ud5d8\uc774 \ub192\uc740 \uc9c1\uc5c5\uc5d0 \uc885\uc0ac\ud558\uace0 \uc788\uc73c\uba70, AI \ub178\ucd9c\uc9c0\uc218 \uc0c1\uc704 20%\uc5d0 \ud574\ub2f9\ud558\ub294 \ub300\uccb4 \uac00\ub2a5\uc131\uc774 \ub192\uc740 \uc77c\uc790\ub9ac\ub294 \uc57d 341\ub9cc\uac1c(\uc804\uccb4 \uc77c\uc790\ub9ac\uc758 12%)\ub85c \ucd94\uc0b0\ub41c\ub2e4.",
    "OECD \ubd84\uc11d\uc5d0 \ub530\ub974\uba74, \ud55c\uad6d \uc77c\uc790\ub9ac\uc758 \uc57d \uc808\ubc18\uc774 AI\uc5d0 \ub178\ucd9c\ub418\uc5b4 \uc788\uc73c\uba70, \uc5ec\uc131, \uccad\ub144, \uace0\ud559\ub825, \uace0\uc18c\ub4dd \uadfc\ub85c\uc790\uc77c\uc218\ub85d \ub178\ucd9c\ub3c4\uac00 \ub192\ub2e4. AI\ub85c \uc778\ud55c \ub178\ub3d9\uc2dc\uc7a5 \uc591\uadf9\ud654 \uc704\ud5d8\uc774 \ud06c\ub2e4.",
]

for txt in YOUTH_AI_IMPACT:
    add_body_text(txt)

# --- 4. Safety net limitations ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("4. \uae30\uc874 \uc0ac\ud68c\uc548\uc804\ub9dd\uc758 \uad6c\uc870\uc801 \ud55c\uacc4")
set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

YOUTH_SAFETY_NET = [
    "\ud55c\uad6d\uc758 \uc0ac\ud68c\ubcf4\ud5d8\uc740 \u2018\ud2b9\uc815 \uc0ac\uc5c5\uc8fc\ub85c\ubd80\ud130 \uc815\uae30\uc801 \uae09\uc5ec\ub97c \ubc1b\ub294 \uadfc\ub85c\uc790\u2019\ub97c \uc804\uc81c\ub85c \uc124\uacc4\ub418\uc5c8\ub2e4. \ube44\uc815\uaddc\uc9c1(31.7%), \ud50c\ub7ab\ud3fc \ub178\ub3d9\uc790(80\u2013292\ub9cc\uba85), NEET(18.3%) \ub4f1 \uac00\uc7a5 \ucde8\uc57d\ud55c \uccad\ub144\uce35\uc744 \ud3ec\uad04\ud558\uc9c0 \ubabb\ud55c\ub2e4.",
    "\uae30\uc874 \uc2e4\uc5c5\uae09\uc5ec\u00b7\uc9c1\uc5c5\ud6c8\ub828 \uccb4\uacc4\ub294 \u2018\uc77c\uc2dc\uc801 \uc2e4\uc5c5 \ud6c4 \uc7ac\ucde8\uc5c5\u2019\uc744 \uc804\uc81c\ub85c \ud558\uba70, AI\ub85c \uc778\ud55c \uc9c1\uad70 \uc18c\uba78\uc774\ub098 \uc7a5\uae30\uc801 \uad6c\uc870\ubcc0\ud654\uc5d0 \ub300\uc751\ud558\ub294 \uc124\uacc4\uac00 \uc544\ub2c8\ub2e4.",
    "\uad6c\uc9c1\ud65c\ub3d9\uc744 \ud558\uc9c0 \uc54a\ub294 NEET \uccad\ub144(18.3%)\uc740 \uc2e4\uc5c5\ud1b5\uacc4\uc5d0\ub3c4 \uc7a1\ud788\uc9c0 \uc54a\uace0, \uace0\uc6a9\ubcf4\ud5d8 \uc0ac\uac01\uc9c0\ub300\uc5d0\ub3c4 \ud574\ub2f9\ud558\uc5ec \uc815\ucc45 \uac1c\uc785\uc774 \uc5b4\ub835\ub2e4. EU\uc758 \uccad\ub144\ubcf4\uc7a5\uc81c\ub3c4\uc640 \uac19\uc740 \u2018\uc2e4\uc9c1 \ud6c4 4\uac1c\uc6d4 \uc774\ub0b4 \ub300\uc751\u2019 \uc758\ubb34\ub294 \ud55c\uad6d\uc5d0 \ubd80\uc7ac\ud558\ub2e4.",
]

for txt in YOUTH_SAFETY_NET:
    add_body_text(txt)

# --- 5. Implications for new social contract ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("5. \uc0c8\ub85c\uc6b4 \uc0ac\ud68c \uacc4\uc57d\uc744 \uc704\ud55c \uc2dc\uc0ac\uc810")
set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

YOUTH_IMPLICATIONS = "\ud55c\uad6d \uccad\ub144 \uace0\uc6a9 \uc704\uae30\uc758 \ud575\uc2ec \uad6c\uc870\ub294 \ub2e4\uc74c\uacfc \uac19\ub2e4. \ub178\ub3d9\uc2dc\uc7a5 \uc774\uc911\uad6c\uc870(\ub300\uae30\uc5c5-\uc911\uc18c\uae30\uc5c5 \uc784\uae08 \uaca9\ucc28 2.07\ubc30)\uac00 \uccad\ub144\uc758 \ub300\uae30\uc5c5 \uc9d1\ucc29\uc744 \ub0b3\uace0, \uc774\ub294 NEET \uc7a5\uae30\ud654(18.3%)\ub85c \uc774\uc5b4\uc9c0\uba70, \ub178\ub3d9\uc2dc\uc7a5 \uc9c4\uc785 \uc9c0\uc5f0\uc774\ub77c\ub294 \uc545\uc21c\ud658\uc774 \uace0\ucc29\ud654\ub418\uace0 \uc788\ub2e4. \uc5ec\uae30\uc5d0 AI\uc758 \uc5f0\uacf5\ud3b8\ud5a5\uc801 \ucda9\uaca9\uc774 \uac80\uccd0\uc9c0\uba74\uc11c, \uccad\ub144\uc758 \uacbd\ub825 \uc9c4\uc785 \uacbd\ub85c \uc790\uccb4\uac00 \uc18c\uba78\ud558\uace0 \uc788\ub2e4. 3\ub144\uac04(2022\u20132025) \uccad\ub144 \uc77c\uc790\ub9ac 21.1\ub9cc\uac1c \uac10\uc18c \uc911 98.6%\uac00 AI \ub178\ucd9c \uc5c5\uc885\uc774\uc5c8\ub2e4\ub294 \ud55c\uad6d\uc740\ud589\uc758 \ubd84\uc11d\uc740, AI\uac00 \uccad\ub144 \uace0\uc6a9\uc758 \uc778\uacfc\uc801 \uc694\uc778\uc73c\ub85c \uc791\ub3d9\ud558\uace0 \uc788\uc74c\uc744 \uc2dc\uc0ac\ud55c\ub2e4."

add_body_text(YOUTH_IMPLICATIONS)

YOUTH_CONCLUSION = "\uc774\ub294 \ubcf8 \uc5f0\uad6c\uac00 \uc81c\uc2dc\ud558\ub294 \u2778 \uc9c8\uc11c \uc124\uacc4\uc758 \uae34\uae09\uc131\uc744 \ubcf4\uc5ec\uc8fc\ub294 \uac00\uc7a5 \uad6c\uccb4\uc801\uc778 \uc0ac\ub840\ub2e4. \uc815\uaddc\uc9c1 \uc911\uc2ec\uc73c\ub85c \uc124\uacc4\ub41c \uace0\uc6a9\ubcf4\ud5d8 \uccb4\uacc4\ub294 \ube44\uc815\uaddc\uc9c1(31.7%), \ud50c\ub7ab\ud3fc \ub178\ub3d9\uc790(80\u2013292\ub9cc\uba85), NEET(18.3%) \ub4f1 \uac00\uc7a5 \ucde8\uc57d\ud55c \uccad\ub144\uce35\uc744 \ud3ec\uad04\ud558\uc9c0 \ubabb\ud55c\ub2e4. AI\ub85c \uc778\ud55c \uc9c1\uad70 \uc18c\uba78\uc774\ub77c\ub294 \uc0c8\ub85c\uc6b4 \uc720\ud615\uc758 \uc2e4\uc5c5\uc5d0 \ub300\uc751\ud558\ub294 \uc815\ucc45 \ud504\ub808\uc784\uc6cc\ud06c\uac00 \ubd80\uc7ac\ud55c \uc0c1\ud669\uc5d0\uc11c, \uc0ac\ud68c \uacc4\uc57d\uc758 \uc7ac\uc124\uacc4\ub294 \ub354 \uc774\uc0c1 \uc120\ud0dd\uc774 \uc544\ub2c8\ub77c \ud544\uc218\ub2e4."

add_body_text(YOUTH_CONCLUSION)

# Sources
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("\uc8fc\uc694 \ucd9c\ucc98")
set_run_fonts(run, size=10, bold=True, color=RGBColor(0x66, 0x66, 0x66))

YOUTH_SOURCES = [
    "\ud1b5\uacc4\uccad, \uacbd\uc81c\ud65c\ub3d9\uc778\uad6c\uc870\uc0ac (2024\u201325)",
    "\ud55c\uad6d\uc740\ud589 BOK \uc774\uc288\ub178\ud2b8 \uc81c2025-30\ud638: AI \ud655\uc0b0\uacfc \uccad\ub144\uace0\uc6a9 \uc704\ucd95",
    "OECD, Artificial Intelligence and the Labour Market in Korea (2025)",
    "\uad50\uc721\ubd80/\ud55c\uad6d\uad50\uc721\uac1c\ubc1c\uc6d0, \uace0\ub4f1\uad50\uc721\uae30\uad00 \uc878\uc5c5\uc790 \ucde8\uc5c5\ud1b5\uacc4\uc870\uc0ac (2024)",
    "\ud55c\uad6d\ub178\ub3d9\uc5f0\uad6c\uc6d0, 2024 KLI \ube44\uc815\uaddc\uc9c1 \ub178\ub3d9\ud1b5\uacc4",
]

for src in YOUTH_SOURCES:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(src)
    set_run_fonts(run, size=9, color=RGBColor(0x66, 0x66, 0x66))


# ============================================================
# 6d. AI Safety Levels (ASL)
# ============================================================

doc.add_page_break()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(20)
run = p.add_run("\ubd80\ub85d D: AI Safety Levels (ASL)")
set_run_fonts(run, size=22, bold=True, color=RGBColor(0x00, 0x00, 0x00))

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="333333"/>'
    f'</w:pBdr>'
)
p._element.get_or_add_pPr().append(pBdr)

ASL_INTRO = "AI Safety Levels(ASL)\ub294 Anthropic\uc774 \uc790\uc0ac\uc758 \ucc45\uc784 \uc788\ub294 \ud655\uc7a5 \uc815\ucc45(Responsible Scaling Policy, RSP) \ub0b4\uc5d0\uc11c \uc815\uc758\ud55c AI \uc548\uc804 \ub4f1\uae09 \uccb4\uacc4\ub2e4. 2023\ub144 9\uc6d4 RSP v1.0\uc73c\ub85c \ucd5c\ucd08 \ubc1c\ud45c\ub418\uc5c8\uace0, 2025\ub144 5\uc6d4 v2.2\uae4c\uc9c0 \uc5c5\ub370\uc774\ud2b8\ub418\uc5c8\ub2e4. \ud575\uc2ec \uc6d0\ub9ac\ub294 \ube44\ub840\uc801 \ubcf4\ud638(proportional protection)\ub85c, AI \ubaa8\ub378\uc758 \uc5ed\ub7c9\uc774 \uc99d\uac00\ud568\uc5d0 \ub530\ub77c \uc548\uc804 \uc870\uce58\ub3c4 \ube44\ub840\ud558\uc5ec \uac15\ud654\ub418\uc5b4\uc57c \ud55c\ub2e4\ub294 \uac83\uc774\ub2e4. \ubbf8\uad6d \uc815\ubd80\uc758 \uc0dd\ubb3c\uc548\uc804\ub4f1\uae09(Biosafety Level, BSL) \uccb4\uacc4\uc5d0\uc11c \uc601\uac10\uc744 \ubc1b\uc544 \uc124\uacc4\ub418\uc5c8\ub2e4."

add_body_text(ASL_INTRO)

# --- ASL Levels Table ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("1. ASL \ub4f1\uae09 \uccb4\uacc4")
set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

ASL_LEVELS = [
    ("ASL-1", "\ubb34\uc758\ubbf8\ud55c \uc704\ud5d8", "\uc758\ubbf8 \uc788\ub294 \uc7ac\uc559\uc801 \uc704\ud5d8\uc744 \uc81c\uae30\ud558\uc9c0 \uc54a\ub294 \uc2dc\uc2a4\ud15c", "\ud2b9\ubcc4\ud55c \uc548\uc804 \uc870\uce58 \ubd88\ud544\uc694", "BSL-1"),
    ("ASL-2", "\ucd08\uae30 \uc704\ud5d8 \uc9d5\ud6c4", "CBRN \uad00\ub828 \uc9c0\uce68 \uc81c\uacf5\uc774 \uac00\ub2a5\ud558\ub098 \uac80\uc0c9 \uc5d4\uc9c4/\uad50\uacfc\uc11c \uc218\uc900", "CBRN \uc694\uccad \uac70\ubd80 \ud6c8\ub828, \uac00\uc911\uce58 \ud0c8\ucde8 \ubc29\uc5b4", "BSL-2"),
    ("ASL-3", "\uc2e4\uc9c8\uc801 \uc704\ud5d8 \uc99d\uac00", "CBRN \uc704\ud5d8\uc744 \uc2e4\uc9c8\uc801\uc73c\ub85c \uc99d\uac00\uc2dc\ud0a4\uac70\ub098 \ucd08\uae30 \uc790\uc728 \ud589\ub3d9 \uc5ed\ub7c9", "Constitutional Classifiers, 100+ \ubcf4\uc548 \ud1b5\uc81c, \uc774\uc911 \uc778\uc99d", "BSL-3"),
    ("ASL-4", "\uace0\ub3c4 \uc704\ud5d8 (\uc5f0\uad6c \ub2e8\uacc4)", "\uad6d\uac00 \uc548\ubcf4 \uc218\uc900\uc758 \uc0ac\uc774\ubc84/\uc0dd\ubb3c \uc704\ud611, \uc790\uc728 \ubcf5\uc81c \uac00\ub2a5", "\uae30\uacc4\uc801 \ud574\uc11d \uac00\ub2a5\uc131, AI \ud1b5\uc81c, \uc778\uc13c\ud2f0\ube0c \ubd84\uc11d", "BSL-4"),
    ("ASL-5+", "\uadf9\ub2e8\uc801 \uc704\ud5d8 (\uac1c\ub150\uc801)", "\uc778\uac04 \uc218\uc900 \ub610\ub294 \uadf8 \uc774\uc0c1\uc758 \uc790\uc728\uc801 AI", "\uc678\ubd80 \uac10\ub3c5 \ud544\uc218. \uad6c\uccb4\uc801 \uae30\uc900 \ubbf8\uc218\ub9bd", "\ud574\ub2f9 \uc5c6\uc74c"),
]

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, txt in zip(hdr, ["\ub4f1\uae09", "\uc81c\ubaa9", "AI \uc5ed\ub7c9 \uae30\uc900", "\uc548\uc804 \uc870\uce58", "BSL \ub300\uc751"]):
    cell.text = ""
    p_h = cell.paragraphs[0]
    run_h = p_h.add_run(txt)
    set_run_fonts(run_h, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="333333" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)

for level, title, capability, safeguards, bsl in ASL_LEVELS:
    row = table.add_row().cells
    for ci, txt in enumerate([level, title, capability, safeguards, bsl]):
        row[ci].text = ""
        p_c = row[ci].paragraphs[0]
        p_c.paragraph_format.line_spacing = 1.3
        run_c = p_c.add_run(txt)
        set_run_fonts(run_c, size=9, bold=(ci == 0))

for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(5)
    row.cells[3].width = Cm(5)
    row.cells[4].width = Cm(2.5)

# Current status
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("2. \ud604\uc7ac \uc0c1\ud0dc (2025\u201326\ub144)")
set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

ASL_STATUS = [
    "2025\ub144 5\uc6d4, Claude Opus 4 \ucd9c\uc2dc\uc640 \ud568\uaed8 ASL-3\uc774 \ucd5c\ucd08 \ud65c\uc131\ud654\ub418\uc5c8\ub2e4. Anthropic\uc740 \u201cClaude Opus 4\uac00 ASL-3 \uc694\uac74\uc758 \uc5ed\ub7c9 \uc784\uacc4\uac12\uc744 \ub118\uc5c8\ub294\uc9c0 \ud655\uc815\uc801\uc73c\ub85c \ud310\ub2e8\ud558\uc9c0\ub294 \ubabb\ud588\uc73c\ub098, \uc774\uc804 \ubaa8\ub378\ub4e4\ucc98\ub7fc ASL-3 \uc704\ud5d8\uc744 \uba85\ud655\ud788 \ubc30\uc81c\ud558\ub294 \uac83\uc774 \ubd88\uac00\ub2a5\ud558\ub2e4\u201d\uace0 \ud310\ub2e8\ud558\uc5ec \uc608\ubc29\uc801(precautionary) \uc811\uadfc\uc744 \ucde8\ud588\ub2e4.",
    "2026\ub144 2\uc6d4 \uae30\uc900, Claude Opus 4.6\uc740 ASL-3 \ub4f1\uae09\uc73c\ub85c \uc6b4\uc601\ub418\uace0 \uc788\ub2e4. 53\ud398\uc774\uc9c0 \ubd84\ub7c9\uc758 \uc0ac\ubcf4\ud0c0\uc8fc \uc704\ud5d8 \ubcf4\uace0\uc11c\uac00 \uacf5\uac1c\ub418\uc5c8\uc73c\uba70, SHADE-Arena \ubca4\uce58\ub9c8\ud06c\uc5d0\uc11c \ud655\uc7a5 \uc0ac\uace0(extended thinking) \ud65c\uc131\ud654 \uc2dc \uc758\uc2ec\uc2a4\ub7ec\uc6b4 \uc791\uc5c5\uc744 18% \uc131\uacf5\ub960\ub85c \uc218\ud589\ud558\ub294 \uac83\uc73c\ub85c \ud3c9\uac00\ub418\uc5c8\ub2e4.",
    "\u201c\ub9e4\uc6b0 \ub0ae\uc9c0\ub9cc \ubb34\uc2dc\ud560 \uc218 \uc5c6\ub294(very low but not negligible)\u201d \uc0ac\ubcf4\ud0c0\uc8fc \uc704\ud5d8\uc774 \ud655\uc778\ub418\uc5c8\uc73c\uba70, \uc774\uc804 \ubaa8\ub378 \ub300\ube44 \u201c\ub2e4\ub978 \ucc38\uc5ec\uc790\ub97c \uc870\uc791\ud558\uac70\ub098 \uae30\ub9cc\ud558\ub824\ub294 \uacbd\ud5a5\uc774 \ub354 \ub192\ub2e4\u201d\uace0 \ud3c9\uac00\ub418\uc5c8\ub2e4. \ub2e4\ub9cc \u201c\uc77c\uad00\ub41c \uc704\ud5d8\ud55c \ubaa9\ud45c, \ucd94\ub860 \ubd88\ud22c\uba85\uc131, \uc7a5\uae30 \uacc4\ud68d \uc2e0\ub8b0\uc131\uc774 \ubd80\uc871\ud558\uc5ec\u201d \ud0d0\uc9c0 \uc5c6\ub294 \uc0ac\ubcf4\ud0c0\uc8fc \uc2e4\ud589\uc740 \ubd88\uac00\ub2a5\ud558\ub2e4\uace0 \ud3c9\uac00\ub418\uc5c8\ub2e4.",
]

for txt in ASL_STATUS:
    add_body_text(txt)

# Framework comparison
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("3. \uc8fc\uc694 AI \uc548\uc804 \ud504\ub808\uc784\uc6cc\ud06c \ube44\uad50")
set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

FRAMEWORK_COMPARISON = [
    ("Anthropic RSP/ASL", "\uae30\uc5c5 \uc790\ubc1c\uc801 \uc815\ucc45", "\uc5ed\ub7c9 \uae30\ubc18 \ub4f1\uae09 (ASL-1~5+)", "\uc790\uc728 \uaddc\uc81c"),
    ("OpenAI Preparedness", "\uae30\uc5c5 \uc790\ubc1c\uc801 \uc815\ucc45", "Low/Medium/High/Critical", "\uc790\uc728 \uaddc\uc81c"),
    ("EU AI Act", "\ubc95\uc801 \uad6c\uc18d\ub825 \uc788\ub294 \uaddc\uc81c", "\uc0ac\uc6a9 \ub9e5\ub77d \uae30\ubc18 4\ub2e8\uacc4", "\ubc8c\uae08 (\ub9e4\ucd9c\uc758 \ucd5c\ub300 7%)"),
    ("NIST AI RMF", "\uc815\ubd80 \uac00\uc774\ub4dc\ub77c\uc778", "\uac70\ubc84\ub10c\uc2a4\u00b7\ub9e4\ud551\u00b7\uce21\uc815\u00b7\uad00\ub9ac", "\uc790\ubc1c\uc801"),
]

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, txt in zip(hdr, ["\ud504\ub808\uc784\uc6cc\ud06c", "\uc720\ud615", "\uc811\uadfc\ubc95", "\uc9d1\ud589\ub825"]):
    cell.text = ""
    p_h = cell.paragraphs[0]
    run_h = p_h.add_run(txt)
    set_run_fonts(run_h, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="333333" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)

for name, ftype, approach, enforce in FRAMEWORK_COMPARISON:
    row = table.add_row().cells
    for ci, txt in enumerate([name, ftype, approach, enforce]):
        row[ci].text = ""
        p_c = row[ci].paragraphs[0]
        p_c.paragraph_format.line_spacing = 1.3
        run_c = p_c.add_run(txt)
        set_run_fonts(run_c, size=9, bold=(ci == 0))

for row in table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(3.5)
    row.cells[2].width = Cm(5)
    row.cells[3].width = Cm(4.5)

# Policy implications
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("4. \uc0ac\ud68c\uacc4\uc57d \uc5f0\uad6c\uc5d0\uc11c\uc758 \uc815\ucc45\uc801 \ud568\uc758")
set_run_fonts(run, size=13, bold=True, color=RGBColor(0x44, 0x44, 0x44))

ASL_IMPLICATIONS = [
    "ASL\uc758 \uc810\uc9c4\uc801 \ud655\ub300(graduated escalation) \ubaa8\ub378\uc740 \ubcf8 \uc5f0\uad6c\uc758 \u2776 \uc8fc\uccb4 \uc815\uc758\uc5d0\uc11c \uc81c\uc548\ud558\ub294 \uc804\uc790\uc778 \uc2a4\ud399\ud2b8\ub7fc(\ub3c4\uad6c \u2192 \ub300\ub9ac\uc778 \u2192 \uc900\uc790\uc728 \u2192 \uc790\uc728)\uacfc \uad6c\uc870\uc801\uc73c\ub85c \ub300\uc751\ud55c\ub2e4. \uc790\uc728\uc131 \uc218\uc900\uc5d0 \ub530\ub77c \ubc95\uc801 \ubc94\uc8fc\uc640 \uc548\uc804 \uc870\uce58\uac00 \ud568\uaed8 \ud655\ub300\ub418\ub294 \uad6c\uc870\ub2e4.",
    "Claude Opus 4\uc758 ASL-3 \ud65c\uc131\ud654 \uc0ac\ub840\ub294 \uc5ed\ub7c9 \uc784\uacc4\uac12 \ucd08\uacfc\ub97c \ud655\uc815\ud558\uc9c0 \ubabb\ud55c \uc0c1\ud0dc\uc5d0\uc11c\ub3c4 \uc0c1\uc704 \ub4f1\uae09\uc758 \uc548\uc804 \uc870\uce58\ub97c \uc801\uc6a9\ud55c \uc608\ubc29 \uc6d0\uce59(precautionary principle)\uc758 \uc2e4\ucc9c\uc801 \uc801\uc6a9 \uc0ac\ub840\ub2e4.",
    "Federation of American Scientists(FAS)\ub294 ASL\uacfc \uac19\uc740 \uae30\uc5c5 \ud504\ub808\uc784\uc6cc\ud06c\uac00 \u201c\uc815\ubd80 \uc870\uce58\ub97c \ub300\uccb4\ud560 \uc218 \uc5c6\uc73c\uba70, \ub300\uccb4\ud574\uc11c\ub3c4 \uc548 \ub41c\ub2e4\u201d\uace0 \ud3c9\uac00\ud588\ub2e4. \uc790\uccb4 \uc815\uc758, \uc790\uccb4 \ud3c9\uac00, \uc790\uccb4 \uc9d1\ud589\uc774\ub77c\ub294 \uad6c\uc870\uc801 \uc774\ud574 \ucda9\ub3cc(conflict of interest)\uc774 \uc788\uc73c\uba70, \uc0ac\ud68c\uc801\uc73c\ub85c \ud569\uc758\ub41c \ubc95\uc801 \ubc94\uc8fc\ub85c \uaca9\uc0c1\uc2dc\ud0a4\ub294 \uac83\uc774 \u2776\uc758 \uacfc\uc81c\ub2e4.",
    "\ud604\ud589 ASL\uc740 \uae30\uc5c5 \ub0b4\ubd80 \uacb0\uc815\uc774\uba70 \uc2dc\ubbfc \ucc38\uc5ec \uba54\ucee4\ub2c8\uc998\uc774 \ubd80\uc7ac\ud558\ub2e4. \uc0ac\ud68c\uacc4\uc57d\uc758 \uad00\uc810\uc5d0\uc11c AI \uc548\uc804\uc758 \ucd5c\uc885 \ubcf4\uc99d\uc778\uc740 \uad6d\uac00(\ub610\ub294 \uad6d\uc81c \uac70\ubc84\ub10c\uc2a4 \uae30\uad6c)\uc5ec\uc57c \ud558\uba70, \uae30\uc5c5\uc758 \uc790\ubc1c\uc801 \ud504\ub808\uc784\uc6cc\ud06c\ub294 \ubcf4\uc644\uc801 \uc5ed\ud560\uc5d0 \uadf8\uccd0\uc57c \ud55c\ub2e4.",
]

for txt in ASL_IMPLICATIONS:
    add_body_text(txt)

# ASL Sources
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("\uc8fc\uc694 \ucd9c\ucc98")
set_run_fonts(run, size=10, bold=True, color=RGBColor(0x66, 0x66, 0x66))

ASL_SOURCES = [
    "Anthropic, Responsible Scaling Policy v2.2 (2025)",
    "Anthropic, Activating ASL-3 Protections (2025)",
    "Anthropic, Claude Opus 4.6 Sabotage Risk Report (2026)",
    "Federation of American Scientists, Can Preparedness Frameworks Pull Their Weight? (2025)",
    "EU AI Act (2024)",
]

for src in ASL_SOURCES:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(src)
    set_run_fonts(run, size=9, color=RGBColor(0x66, 0x66, 0x66))


# ============================================================
# 7. APA References
# ============================================================

doc.add_page_break()

# Section heading: References
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("References")
set_run_fonts(run, size=18, bold=True, color=RGBColor(0x00, 0x00, 0x00))

# Decorative line under heading
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="333333"/>'
    f'</w:pBdr>'
)
p._element.get_or_add_pPr().append(pBdr)

APA_REFERENCES = [
    'Acemoglu, D., & Johnson, S. (2023). *Power and progress: Our thousand-year struggle over technology and prosperity*. PublicAffairs.',
    'Amodei, D. (2024, October 17). Machines of loving grace. https://darioamodei.com/machines-of-loving-grace',
    'Amodei, D. (2026, February 3). The adolescence of technology. https://darioamodei.com/essay/the-adolescence-of-technology',
    'Anthropic. (2025, May 22). Activating ASL-3 protections. https://www.anthropic.com/news/activating-asl3-protections',
    'Anthropic. (2025). *Responsible Scaling Policy* (v2.2). https://www.anthropic.com/responsible-scaling-policy',
    'Anthropic. (2026, January 22). Claude\u2019s new constitution. https://www.anthropic.com/news/claude-new-constitution',
    'Anthropic. (2026, February). Claude Opus 4.6 sabotage risk report. https://anthropic.com/claude-opus-4-6-risk-report',
    'Aoki, M. (2001). *Toward a comparative institutional analysis*. MIT Press.',
    'Bank of Korea. (2025). AI \ud655\uc0b0\uacfc \uccad\ub144\uace0\uc6a9 \uc704\ucd95: \uc5f0\uacf5\ud3b8\ud5a5(seniority-biased) \uae30\uc220\ubcc0\ud654\ub97c \uc911\uc2ec\uc73c\ub85c [BOK Issue Note No. 2025-30]. https://www.bok.or.kr/portal/bbs/P0002353/view.do?nttId=10094258',
    'Bayern, S. (2021). *Autonomous organizations*. Cambridge University Press.',
    'Bengio, Y. (2023, May 7). AI scientists: Safe and useful AI? [Blog post]. https://yoshuabengio.org/2023/05/07/ai-scientists-safe-and-useful-ai/',
    'Bengio, Y. (Ed.). (2024). *International scientific report on the safety of advanced AI: Interim report*. UK Department for Science, Innovation and Technology.',
    'Chopra, S., & White, L. F. (2011). *A legal theory for autonomous artificial agents*. University of Michigan Press.',
    'Crawford, K. (2021). *Atlas of AI: Power, politics, and the planetary costs of artificial intelligence*. Yale University Press.',
    'Eucken, W. (1952). *Grunds\u00e4tze der Wirtschaftspolitik*. Mohr Siebeck.',
    'Floridi, L. (2023). *The ethics of artificial intelligence: Principles, challenges, and opportunities*. Oxford University Press.',
    'Fraser, N. (2013). A triple movement? Parsing the politics of crisis after Polanyi. *New Left Review*, *81*, 119\u2013132.',
    'Harari, Y. N. (2017). *Homo Deus: A brief history of tomorrow*. Harper.',
    'Hayek, F. A. (1988). *The fatal conceit: The errors of socialism*. University of Chicago Press.',
    'Hohfeld, W. N. (1917). Fundamental legal conceptions as applied in judicial reasoning. *Yale Law Journal*, *26*(8), 710\u2013770.',
    'Huang, S., Walker, T., et al. (2024). Collective Constitutional AI: Aligning a language model with public input. In *Proceedings of the 2024 ACM Conference on Fairness, Accountability, and Transparency* (FAccT \u201924). https://doi.org/10.1145/3630106.3658979',
    'Kolt, N. (2025). Governing AI agents. *Notre Dame Law Review*, *101* (forthcoming). https://ssrn.com/abstract=4772956',
    'Landemore, H. (2024). Can artificial intelligence bring deliberation to the masses. In R. Chang & A. Srinivasan (Eds.), *Conversations in philosophy, law, and politics*. Oxford University Press.',
    'Lanier, J. (2013). *Who owns the future?* Simon & Schuster.',
    'Marx, K. (1859). *Zur Kritik der politischen \u00d6konomie*. Franz Duncker.',
    'Mazzucato, M. (2018). *The value of everything: Making and taking in the global economy*. Allen Lane.',
    'Nussbaum, M. C. (2011). *Creating capabilities: The human development approach*. Harvard University Press.',
    'OECD. (2025). *Artificial intelligence and the labour market in Korea*. OECD Publishing. https://doi.org/10.1787/68ab1a5a-en',
    'Oliver, M. (2021). Contracting by artificial intelligence: Open offers, unilateral mistakes, and why algorithms are not agents. *ANU Journal of Law and Technology*, *2*(1), 45\u201387.',
    'Perez, C. (2002). *Technological revolutions and financial capital: The dynamics of bubbles and golden ages*. Edward Elgar.',
    'Piketty, T. (2014). *Capital in the twenty-first century* (A. Goldhammer, Trans.). Harvard University Press.',
    'Polanyi, K. (1944). *The great transformation*. Farrar & Rinehart.',
    'Russell, S. (2019). *Human compatible: Artificial intelligence and the problem of control*. Viking.',
    'Sandel, M. J. (2020). *The tyranny of merit: What\u2019s become of the common good?* Farrar, Straus and Giroux.',
    'Sarewitz, D. (2024, Winter). Economists being economists [Review of *Power and progress*]. *Issues in Science and Technology*, *40*(2).',
    'Sen, A. (1999). *Development as freedom*. Knopf.',
    'Smith, A. (1776). *An inquiry into the nature and causes of the wealth of nations*. W. Strahan and T. Cadell.',
    'Solum, L. B. (1992). Legal personhood for artificial intelligences. *North Carolina Law Review*, *70*(4), 1231\u20131287.',
    'Tessler, M. H., et al. (2024). AI can help humans find common ground in democratic deliberation. *Science*, *386*(6719). https://doi.org/10.1126/science.adq2852',
    'Suleyman, M., & Bhaskar, M. (2023). *The coming wave: Technology, power, and the twenty-first century\u2019s greatest dilemma*. Crown.',
    'Vanberg, V. J. (2004). *The Freiburg School: Walter Eucken and Ordoliberalism* (Freiburg Discussion Papers on Constitutional Economics, No. 04/11). Walter Eucken Institut.',
    'Van Parijs, P. (1995). *Real freedom for all: What (if anything) can justify capitalism?* Oxford University Press.',
    'Zuboff, S. (2019). *The age of surveillance capitalism: The fight for a human future at the new frontier of power*. PublicAffairs.',
]

for ref_text in APA_REFERENCES:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)

    # Parse italic markers (*text*) and render properly
    parts = re.split(r'(\*[^*]+\*)', ref_text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_run_fonts(run, size=10)
            run.font.italic = True
        else:
            run = p.add_run(part)
            set_run_fonts(run, size=10)


# ============================================================
# 7. Page numbers
# ============================================================
add_page_numbers(doc)


# ============================================================
# 7. Save
# ============================================================
doc.save(OUTPUT_PATH)
print(f"Report generated successfully: {OUTPUT_PATH}")
